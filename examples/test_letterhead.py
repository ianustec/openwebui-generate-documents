"""Local smoke tests for letterhead helpers (no Open WebUI required)."""
from __future__ import annotations

import asyncio
import importlib.util
import sys
import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

BASE = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BASE))

_spec = importlib.util.spec_from_file_location(
    "generate_documents", BASE / "generate_documents.py"
)
mod = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(mod)


def _make_letterhead_bytes() -> bytes:
    doc = Document()
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    hp.text = "ACME Letterhead HEADER"
    fp = section.footer.paragraphs[0]
    fp.text = "ACME Letterhead FOOTER"
    body = doc.add_paragraph("OLD BODY TEXT THAT MUST DISAPPEAR")
    body.runs[0].font.size = Pt(11)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_match_rules():
    cands = [
        {"id": "1", "name": "Carta_Rossi.docx"},
        {"id": "2", "name": "Quote_Sample.DOCX"},
    ]
    assert mod._match_letterhead("Carta_Rossi.docx", cands)["id"] == "1"
    assert mod._match_letterhead("carta_rossi.docx", cands)["id"] == "1"
    assert mod._match_letterhead("Carta_Rossi", cands)["id"] == "1"
    assert mod._match_letterhead("Quote_Sample", cands)["id"] == "2"
    assert mod._match_letterhead("missing.docx", cands) is None
    err = mod._letterhead_missing_error("missing.docx", cands)
    assert "Carta_Rossi.docx" in err and "Quote_Sample.DOCX" in err
    print("OK match rules")


def test_chat_docx_files_filter():
    meta = {
        "files": [
            {"id": "a", "filename": "notes.txt"},
            {"id": "b", "name": "Carta.docx"},
            {"file": {"id": "c", "filename": "Other.dotx"}},
            {"id": "d", "filename": "photo.png"},
        ]
    }
    out = mod._chat_docx_files(meta, None)
    names = {x["name"] for x in out}
    assert names == {"Carta.docx", "Other.dotx"}
    print("OK chat file filter")


def test_clear_body_preserves_header():
    raw = _make_letterhead_bytes()
    doc = mod._load_letterhead_doc(raw=raw)
    assert "OLD BODY" in doc.paragraphs[0].text
    assert doc.sections[0].header.paragraphs[0].text == "ACME Letterhead HEADER"
    mod._clear_body(doc)
    # Body should be empty of paragraphs (sectPr only); header remains.
    body = doc.element.body
    assert body.find(qn("w:sectPr")) is not None
    assert all(child.tag == qn("w:sectPr") for child in body)
    assert doc.sections[0].header.paragraphs[0].text == "ACME Letterhead HEADER"
    assert doc.sections[0].footer.paragraphs[0].text == "ACME Letterhead FOOTER"
    print("OK clear_body preserves header/footer")


async def test_build_with_letterhead_bytes():
    raw = _make_letterhead_bytes()
    tool = mod.Tools()
    spec = mod._resolve_template({
        "template": "blank",
        "letterhead": "Carta_Rossi.docx",
        "title": "Test",
        "_letterhead_bytes": raw,
        "blocks": [
            {"type": "heading", "level": 1, "text": "New content"},
            {"type": "paragraph", "text": "Generated body line."},
        ],
    })
    doc = await tool._build_document(spec)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("New content" in t for t in texts)
    assert any("Generated body" in t for t in texts)
    assert not any("OLD BODY" in t for t in texts)
    assert doc.sections[0].header.paragraphs[0].text == "ACME Letterhead HEADER"
    print("OK build with letterhead bytes")


async def test_build_without_letterhead_regression():
    tool = mod.Tools()
    source = (BASE / "examples" / "report.md").read_text()
    raw = tool._parse_content(source)
    raw.setdefault("template", "report")
    spec = mod._resolve_template(raw)
    assert not spec.get("letterhead")
    doc = await tool._build_document(spec)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as fh:
        doc.save(fh.name)
        size = Path(fh.name).stat().st_size
    assert size > 5000
    print("OK no-letterhead regression", size)


def test_coalesce_aliases():
    s = mod._resolve_template({"template": "blank", "sample": "X.docx"})
    assert s.get("letterhead") == "X.docx"
    s2 = mod._resolve_template({"template": "blank", "base_docx": "Y.docx"})
    assert s2.get("letterhead") == "Y.docx"
    print("OK coalesce aliases")


def main():
    test_match_rules()
    test_chat_docx_files_filter()
    test_clear_body_preserves_header()
    test_coalesce_aliases()
    asyncio.run(test_build_with_letterhead_bytes())
    asyncio.run(test_build_without_letterhead_regression())
    print("ALL SMOKE TESTS PASSED")


if __name__ == "__main__":
    main()
