"""
title: Generate Documents
author: IANUSTEC
author_url: https://ianustec.com
funding_url: https://github.com/ianustec
description: Generate high-quality native Word (.docx) documents from Markdown or a JSON spec - cover pages, styled headings, tables, callouts, TOC, header/footer
requirements: python-docx, Pillow, httpx, pydantic, lxml, markdown-it-py, mdit-py-plugins, PyYAML
required_open_webui_version: 0.4.0
version: 1.1.0
license: MIT
"""

# ============================================================================
# Native DOCX Documents Engine — open-source OpenWebUI tool
# ----------------------------------------------------------------------------
# Server-side OpenWebUI tool `generate_document`. The model emits Markdown
# with YAML frontmatter (preferred) or a structured JSON spec, and this engine
# renders a NATIVE .docx with python-docx: a coherent design system, cover
# pages, numbered headings with rules, styled tables, callouts, code blocks,
# signatures, TOC and running header/footer.
#
# DUAL INPUT (auto-detected):
#   1. Markdown with YAML frontmatter (preferred — compact, robust on long docs):
#
#          ---
#          template: report
#          title: FY 2024 Report
#          cover: auto
#          ---
#
#          # Executive summary
#          Revenue grew **24%** ...
#
#          ::: callout type="success" title="Highlight"
#          Margin expanded by ==6 pp==.
#          :::
#
#   2. Structured JSON (legacy, still supported): object with fields
#      title/template/page/styles/header/footer/cover/blocks.
#
# Both formats converge on the same dict shape, then flow through template
# merge, OOXML rendering, the (optional) image pipeline and save.
#
# The document is saved via the OpenWebUI Files API (with a /cache/files
# fallback) and a clickable download link is emitted in chat.
#
# License: MIT — Copyright (c) IANUSTEC.
# ============================================================================
# region ── Imports ────────────────────────────────────────────────────────────

from __future__ import annotations

import base64
import json
import os
import re
import sys
import traceback
import unicodedata
import uuid
from io import BytesIO
from typing import Any, Optional

import httpx
from pydantic import BaseModel, Field

try:
    from PIL import Image  # type: ignore
except ImportError:
    Image = None  # type: ignore[assignment]

# Markdown parser stack (for the dual-input pipeline). Wrapped in try/except
# so the module still imports if these deps are missing — in that case the
# JSON path keeps working and the markdown path returns a clear error.
try:
    import yaml  # type: ignore
    from markdown_it import MarkdownIt  # type: ignore
    from mdit_py_plugins.front_matter import front_matter_plugin  # type: ignore
    from mdit_py_plugins.container import container_plugin  # type: ignore
    _HAS_MD_PARSER = True
except ImportError:
    yaml = None  # type: ignore[assignment]
    MarkdownIt = None  # type: ignore[assignment]
    front_matter_plugin = None  # type: ignore[assignment]
    container_plugin = None  # type: ignore[assignment]
    _HAS_MD_PARSER = False

# python-docx is the engine. We import the full set of low-level building
# blocks once at module load so the renderers can stay terse.
from docx import Document  # type: ignore
from docx.document import Document as _DocxDocument  # type: ignore
from docx.enum.section import WD_ORIENTATION  # type: ignore
from docx.enum.style import WD_STYLE_TYPE  # type: ignore
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE  # type: ignore
from docx.enum.text import (  # type: ignore
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_TAB_ALIGNMENT,
)
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Mm, Pt, RGBColor, Emu  # type: ignore

# Optional: OpenWebUI Files API integration (only available inside the
# OpenWebUI runtime). When unavailable, the tool still works via a local
# /cache/files fallback (download-only, no UI preview).
try:
    from fastapi import UploadFile  # type: ignore
    from starlette.datastructures import Headers  # type: ignore
    from open_webui.routers.files import upload_file_handler  # type: ignore
    from open_webui.models.users import Users  # type: ignore
    _HAS_OWUI_FILES = True
except ImportError:
    _HAS_OWUI_FILES = False

try:
    from open_webui.routers.images import image_generations as _owui_image_generations  # type: ignore
    from open_webui.routers.images import CreateImageForm as _OwuiCreateImageForm  # type: ignore
    from open_webui.models.users import UserModel as _OwuiUserModel  # type: ignore
    from open_webui.models.files import Files as _OwuiFiles  # type: ignore
    _HAS_OWUI_IMAGES = True
except ImportError:
    _HAS_OWUI_IMAGES = False

# Optional storage abstraction so we can read raw bytes back regardless of
# whether OpenWebUI is configured with local disk, S3/MinIO, GCS or Azure.
# When the path stored in the ``file`` table is something like
# ``s3://bucket/key`` it won't exist on local disk; ``Storage.get_file()``
# downloads the object to ``UPLOAD_DIR`` and returns a local path.
try:
    from open_webui.storage.provider import Storage as _OwuiStorage  # type: ignore
    _HAS_OWUI_STORAGE = True
except ImportError:
    _OwuiStorage = None  # type: ignore[assignment]
    _HAS_OWUI_STORAGE = False

# endregion


# region ── Constants & defaults ───────────────────────────────────────────────

# DXA = "DistanceXAxis" = twentieths of a point. python-docx uses Emu/Pt/Mm/Cm
# wrappers but some low-level operations (column widths, page size when bypass
# the Section helpers) need raw DXA. 1 DXA = 1/20 pt = 1/1440 inch.
DXA_PER_INCH = 1440
EMU_PER_INCH = 914400  # used by ImageRun (python-docx wraps it in Emu())

# Default body font. Arial is universally supported and ships with every
# Office/LibreOffice install. The skill's recommendation.
DEFAULT_FONT = "Arial"
DEFAULT_FONT_SIZE_PT = 11
DEFAULT_LINE_SPACING = 1.15

# Anthropic-grade typography defaults (tune cautiously — these are global).
# A refined, tighter scale: H1 anchors sections, H2/H3/H4 step down gently.
HEADING_SIZES_PT = {1: 16.5, 2: 12.5, 3: 11.5, 4: 10.5}
HEADING_BEFORE_PT = {1: 14, 2: 9, 3: 7, 4: 5}
HEADING_AFTER_PT = {1: 6, 2: 3, 3: 2, 4: 2}
# Levels that get an automatic hairline rule beneath them (section dividers).
# Only H1 anchors a section with a divider rule; sub-headings stay clean.
HEADING_RULE_LEVELS = (1,)
# H1/H2 kickers use a touch of tracking; H4 is uppercased for an eyebrow feel.
HEADING_UPPERCASE_LEVELS: tuple[int, ...] = ()

# Default callout palettes (background fill, border accent). Used by the
# "callout" block type.
CALLOUT_PALETTES: dict[str, dict[str, str]] = {
    "info": {"fill": "EAF3FB", "border": "2E75B6", "icon_text": "i"},
    "warning": {"fill": "FFF4E5", "border": "ED7D31", "icon_text": "!"},
    "note": {"fill": "F2F2F2", "border": "7F7F7F", "icon_text": "•"},
    "success": {"fill": "E8F5E9", "border": "4CAF50", "icon_text": "✓"},
    "danger": {"fill": "FDECEA", "border": "D32F2F", "icon_text": "×"},
}

# Page size table (Mm). Anthropic skill recommends explicit page setup —
# python-docx defaults to Letter. We default to A4 (Italian/EU standard).
PAGE_SIZES_MM: dict[str, tuple[float, float]] = {
    "a4": (210.0, 297.0),
    "letter": (215.9, 279.4),
    "legal": (215.9, 355.6),
    "a3": (297.0, 420.0),
    "a5": (148.0, 210.0),
}
DEFAULT_PAGE_SIZE = "a4"
DEFAULT_MARGIN_MM = 25.0  # 25 mm = ~1 inch, comfortable B2B default

# Smart quote regex (same logic as openwebui_artifacts._smart_quotes, kept
# in-module so this tool has zero cross-imports from the slides tool).

# endregion


# region ── Image pipeline (mirrors openwebui_artifacts patterns) ─────────────

async def _fetch_unsplash_image(
    query: str,
    access_key: str,
    *,
    timeout: int = 15,
    orientation: str = "landscape",
) -> Optional[bytes]:
    """Search Unsplash and return the regular-size JPEG bytes, or None."""
    if not access_key or not query:
        return None
    try:
        async with httpx.AsyncClient(timeout=timeout) as client:
            resp = await client.get(
                "https://api.unsplash.com/search/photos",
                params={"query": query, "per_page": 1, "orientation": orientation},
                headers={"Authorization": f"Client-ID {access_key}"},
            )
            if resp.status_code != 200:
                return None
            results = resp.json().get("results", [])
            if not results:
                return None
            img_url = results[0].get("urls", {}).get("regular")
            if not img_url:
                return None
            img_resp = await client.get(img_url)
            if img_resp.status_code == 200:
                return img_resp.content
    except Exception:
        traceback.print_exc()
    return None


async def _generate_ai_image(
    prompt: str,
    api_url: str,
    api_key: str,
    *,
    timeout: int = 60,
) -> Optional[bytes]:
    """Generate an image via OpenAI-compatible /v1/images/generations."""
    if not api_url or not prompt:
        return None
    try:
        headers: dict[str, str] = {"Content-Type": "application/json"}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        body = {
            "prompt": prompt,
            "n": 1,
            "size": "1024x1024",
            "response_format": "url",
        }
        async with httpx.AsyncClient(timeout=timeout) as client:
            resp = await client.post(
                f"{api_url.rstrip('/')}/images/generations",
                headers=headers,
                json=body,
            )
            if resp.status_code != 200:
                return None
            data = resp.json().get("data", [])
            if not data:
                return None
            img_url = data[0].get("url") or data[0].get("b64_json")
            if not img_url:
                return None
            if img_url.startswith("data:") or len(img_url) > 2000:
                return base64.b64decode(data[0].get("b64_json", ""))
            img_resp = await client.get(img_url)
            if img_resp.status_code == 200:
                return img_resp.content
    except Exception:
        traceback.print_exc()
    return None


def _decode_image_b64(value: str) -> Optional[bytes]:
    """Decode an inline base64 image. Accepts both raw b64 and data URIs."""
    if not value or not isinstance(value, str):
        return None
    s = value.strip()
    if s.startswith("data:"):
        # data:image/png;base64,XXXX
        try:
            s = s.split(",", 1)[1]
        except IndexError:
            return None
    try:
        return base64.b64decode(s, validate=False)
    except Exception:
        return None


def _normalise_image_bytes(raw: bytes, *, max_width_px: int = 1600) -> tuple[bytes, str]:
    """Normalise an image to JPEG/PNG bytes + extension hint.

    python-docx's ``add_picture`` accepts a file-like object and infers the
    type from the byte signature, so we don't strictly need to convert. But
    Pillow normalisation lets us:
      - downscale very large images (LLMs sometimes hand back 4096×4096)
      - convert exotic modes (RGBA on a JPEG path) to safe RGB
      - strip alpha when going through JPEG to avoid black backgrounds
    Returns (bytes, extension_without_dot).
    """
    if Image is None:
        # No Pillow → trust the caller, assume PNG-safe
        return raw, "png"
    try:
        img = Image.open(BytesIO(raw))
        img.load()
        orig_format = (img.format or "PNG").upper()
        # Downscale if too wide (preserves aspect ratio)
        if img.width > max_width_px:
            ratio = max_width_px / img.width
            new_size = (max_width_px, max(1, int(img.height * ratio)))
            img = img.resize(new_size, Image.LANCZOS)
        # Mode normalisation: PNG keeps alpha, JPEG drops it
        if orig_format == "PNG":
            buf = BytesIO()
            img.save(buf, format="PNG", optimize=True)
            return buf.getvalue(), "png"
        # Default: JPEG (good for photos, smaller files)
        if img.mode in ("RGBA", "LA", "P"):
            img = img.convert("RGB")
        buf = BytesIO()
        img.save(buf, format="JPEG", quality=88, optimize=True)
        return buf.getvalue(), "jpg"
    except Exception:
        traceback.print_exc()
        return raw, "png"


async def _fetch_image_url(url: str, *, timeout: int = 15) -> Optional[bytes]:
    """Best-effort GET on a public image URL. Returns None on any failure.

    Used by the markdown path to support ``![alt](https://...)`` directly,
    so the model can drop in arbitrary image URLs without going through the
    Unsplash / AI-gen resolvers.
    """
    if not url or not url.lower().startswith(("http://", "https://")):
        return None
    try:
        async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
            r = await client.get(url)
            r.raise_for_status()
            ctype = (r.headers.get("content-type") or "").lower()
            if not ctype.startswith("image/"):
                return None
            return r.content
    except Exception:
        return None


def _b64_decode_lenient(data: str) -> Optional[bytes]:
    """Decode a base64 payload tolerating missing padding and ``data:`` prefixes.

    Mirrors ``openwebui_artifacts._b64_decode_lenient``; defensive against the
    ComfyUI-via-llm-neura ``b64_json`` payloads whose length isn't a multiple
    of 4 (which makes strict ``base64.b64decode`` raise ``Incorrect padding``).
    """
    if not data or not isinstance(data, str):
        return None
    s = data.strip()
    if s.startswith("data:"):
        _, _, s = s.partition(",")
    s = re.sub(r"\s+", "", s)
    s += "=" * (-len(s) % 4)
    try:
        return base64.b64decode(s, validate=False)
    except Exception:
        return None


def _read_owui_file_bytes(file_id: str) -> Optional[bytes]:
    """Read raw bytes for an OpenWebUI file, supporting any storage backend.

    Mirrors ``openwebui_artifacts._read_owui_file_bytes``. The path stored on
    the ``file`` row may be a local path or an opaque URI (``s3://``, ``gs://``,
    ``https://...``); we always try local first and fall back to the
    ``Storage`` abstraction which downloads remote objects to ``UPLOAD_DIR``.
    The cached copy is removed after read so we don't accumulate temp files.
    """
    if not _HAS_OWUI_IMAGES or not file_id:
        return None
    try:
        file_row = _OwuiFiles.get_file_by_id(file_id)
    except Exception as exc:
        print(f"[documents] image: file lookup failed for {file_id}: {exc}", file=sys.stderr)
        return None
    if not file_row or not getattr(file_row, "path", None):
        print(f"[documents] image: no file row / path for id {file_id}", file=sys.stderr)
        return None
    path = file_row.path
    if isinstance(path, str) and os.path.isfile(path):
        try:
            with open(path, "rb") as fh:
                return fh.read()
        except Exception as exc:
            print(f"[documents] image: local read failed for {path}: {exc}", file=sys.stderr)
    if not _HAS_OWUI_STORAGE or _OwuiStorage is None:
        print(
            f"[documents] image: path not local and Storage unavailable: {path}",
            file=sys.stderr,
        )
        return None
    cached_path: Optional[str] = None
    try:
        cached_path = _OwuiStorage.get_file(path)
        if not cached_path or not os.path.isfile(cached_path):
            print(
                f"[documents] image: Storage.get_file returned no local path for {path}",
                file=sys.stderr,
            )
            return None
        with open(cached_path, "rb") as fh:
            return fh.read()
    except Exception as exc:
        print(f"[documents] image: Storage download failed for {path}: {exc}", file=sys.stderr)
        return None
    finally:
        if (
            cached_path
            and isinstance(path, str)
            and not path.startswith(("/", "."))
            and os.path.isfile(cached_path)
        ):
            try:
                os.remove(cached_path)
            except Exception:
                pass


async def _direct_image_generation(prompt: str, request: Any) -> Optional[bytes]:
    """Bypass ``image_generations`` and call the configured OpenAI-compatible
    upstream directly. Used as a recovery path when OpenWebUI's own route
    crashes (e.g. the ``Incorrect padding`` chain triggered by some upstreams)
    so the document still gets its image.
    """
    cfg = getattr(getattr(request, "app", None), "state", None)
    cfg = getattr(cfg, "config", None) if cfg else None
    if cfg is None:
        return None
    engine = (getattr(cfg, "IMAGE_GENERATION_ENGINE", "") or "").lower()
    base_url = (
        getattr(cfg, "IMAGES_OPENAI_API_BASE_URL", "")
        or getattr(cfg, "IMAGE_OPENAI_API_BASE_URL", "")
        or ""
    ).rstrip("/")
    api_key = (
        getattr(cfg, "IMAGES_OPENAI_API_KEY", "")
        or getattr(cfg, "IMAGE_OPENAI_API_KEY", "")
        or ""
    )
    model = getattr(cfg, "IMAGE_GENERATION_MODEL", "") or ""
    size = getattr(cfg, "IMAGE_SIZE", "1024x1024") or "1024x1024"
    if engine and engine not in ("", "openai"):
        return None
    if not base_url:
        return None
    body: dict[str, Any] = {
        "prompt": prompt,
        "n": 1,
        "size": size,
        "response_format": "b64_json",
    }
    if model:
        body["model"] = model
    headers = {"Content-Type": "application/json"}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    try:
        async with httpx.AsyncClient(timeout=120) as client:
            resp = await client.post(f"{base_url}/images/generations", headers=headers, json=body)
            if resp.status_code != 200:
                print(
                    f"[documents] image: direct upstream HTTP {resp.status_code}",
                    file=sys.stderr,
                )
                return None
            payload = resp.json().get("data") or []
            if not payload:
                return None
            entry = payload[0]
            b64 = entry.get("b64_json")
            if b64:
                data = _b64_decode_lenient(b64)
                if data:
                    return data
            img_url = entry.get("url")
            if img_url:
                if img_url.startswith("data:"):
                    _, _, body64 = img_url.partition(",")
                    return _b64_decode_lenient(body64)
                img_resp = await client.get(img_url)
                if img_resp.status_code == 200:
                    return img_resp.content
    except Exception as exc:
        print(f"[documents] image: direct upstream call failed: {exc}", file=sys.stderr)
    return None


async def _generate_image_via_openwebui(
    prompt: str,
    request: Any,
    user_dict: Optional[dict],
) -> Optional[bytes]:
    """Generate an image via OpenWebUI's internal image router.

    Mirrors the helper in ``openwebui_artifacts.py``. The router persists the
    PNG as a regular OpenWebUI file; we read the bytes via the ``Storage``
    abstraction so any backend (local, S3/MinIO, GCS, Azure) is supported,
    and we embed them inline in the docx so the resulting file is fully
    self-contained — required because the user downloads the docx and any
    ``s3://...`` URL would be unreachable from their machine.
    """
    if not _HAS_OWUI_IMAGES:
        print("[documents] image: openwebui image router unavailable", file=sys.stderr)
        return None
    if not (request and user_dict and prompt):
        print(
            f"[documents] image: skip openwebui (request={bool(request)} "
            f"user={bool(user_dict)} prompt={bool(prompt)})",
            file=sys.stderr,
        )
        return None
    try:
        user_model = (
            _OwuiUserModel(**user_dict) if isinstance(user_dict, dict) else user_dict
        )
    except Exception as exc:
        print(f"[documents] image: cannot build UserModel: {exc}", file=sys.stderr)
        return None
    try:
        result = await _owui_image_generations(
            request=request,
            form_data=_OwuiCreateImageForm(prompt=prompt),
            user=user_model,
        )
    except Exception as exc:
        msg = str(exc).lower()
        is_padding_bug = (
            "padding" in msg
            or ("nonetype" in msg and "lower" in msg)
            or "guess_extension" in msg
        )
        if is_padding_bug:
            print(
                "[documents] image: openwebui image_generations crashed on b64 padding "
                "— trying direct upstream fallback",
                file=sys.stderr,
            )
            direct = await _direct_image_generation(prompt, request)
            if direct:
                return direct
        else:
            print(f"[documents] image: openwebui image_generations failed: {exc}", file=sys.stderr)
            traceback.print_exc()
        return None
    if not result:
        print("[documents] image: openwebui returned empty result", file=sys.stderr)
        return None
    first = result[0] if isinstance(result, list) else result
    url = (first.get("url") if isinstance(first, dict) else None) or (
        first if isinstance(first, str) else None
    )
    if not url:
        print(f"[documents] image: openwebui result has no url: {first!r}", file=sys.stderr)
        return None

    if url.startswith("data:image"):
        _, _, payload = url.partition(",")
        return _b64_decode_lenient(payload) if payload else None

    file_id_match = re.search(r"/files/([0-9a-fA-F-]{8,})/", url)
    if file_id_match:
        data = _read_owui_file_bytes(file_id_match.group(1))
        if data:
            return data
        print(
            f"[documents] image: could not read bytes for {url} — trying direct upstream",
            file=sys.stderr,
        )
        direct = await _direct_image_generation(prompt, request)
        if direct:
            return direct
        return None

    if url.startswith("/"):
        try:
            async with httpx.AsyncClient(timeout=30) as client:
                resp = await client.get(f"http://127.0.0.1:8080{url}")
                if resp.status_code == 200:
                    return resp.content
                print(
                    f"[documents] image: localhost fetch {url} → HTTP {resp.status_code}",
                    file=sys.stderr,
                )
        except Exception as exc:
            print(f"[documents] image: localhost fetch failed: {exc}", file=sys.stderr)
        return None

    try:
        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.get(url)
            if resp.status_code == 200:
                return resp.content
            print(
                f"[documents] image: remote fetch {url[:80]}... → HTTP {resp.status_code}",
                file=sys.stderr,
            )
    except Exception as exc:
        print(f"[documents] image: remote fetch failed: {exc}", file=sys.stderr)
    return None


async def _resolve_image_bytes(
    block: dict,
    *,
    unsplash_key: str,
    ai_url: str,
    ai_key: str,
    request: Any = None,
    user_dict: Optional[dict] = None,
) -> Optional[bytes]:
    """Resolve an image block to raw bytes. Tries in priority order:

    1. ``image_b64`` (or ``image_data``) — explicit base64 (or data URI)
    2. ``image_generate`` — AI prompt:
       a) OpenWebUI internal image router (uses the user's configured
          chat image backend — no extra config needed)
       b) explicit ``ai_url`` valve as fallback
    3. ``image_hint`` — Unsplash search query (only if access key configured)
    4. ``image_url`` — direct HTTPS URL (markdown ``![alt](https://...)``)

    Returns ``None`` if none of the strategies succeed; the caller should
    silently skip the image block (mirrors the slides tool behaviour).
    """
    raw_b64 = block.get("image_b64") or block.get("image_data")
    if raw_b64:
        decoded = _decode_image_b64(raw_b64)
        if decoded:
            return decoded
    if block.get("image_generate"):
        gen = await _generate_image_via_openwebui(
            block["image_generate"], request, user_dict
        )
        if gen is None and ai_url:
            gen = await _generate_ai_image(block["image_generate"], ai_url, ai_key)
        if gen:
            return gen
    if block.get("image_hint"):
        if unsplash_key:
            unsp = await _fetch_unsplash_image(
                block["image_hint"],
                unsplash_key,
                orientation=block.get("orientation", "landscape"),
            )
            if unsp:
                return unsp
        # Fall back to AI generation from the hint when Unsplash is missing.
        gen = await _generate_image_via_openwebui(
            block["image_hint"], request, user_dict
        )
        if gen is None and ai_url:
            gen = await _generate_ai_image(block["image_hint"], ai_url, ai_key)
        if gen:
            return gen
    if block.get("image_url"):
        url_bytes = await _fetch_image_url(block["image_url"])
        if url_bytes:
            return url_bytes
    return None

# endregion


# region ── Text helpers ──────────────────────────────────────────────────────

def _smart_quotes(text: str) -> str:
    """Replace ASCII quotes with curly quotes for typographic polish.

    - "..." → "..."
    - apostrophe ' between letters → ' (l'azienda, don't, it's)
    - " - " (spaced hyphen) → em-dash " — "
    Idempotent: skips if curly quotes already present in the input.
    """
    if not text or not isinstance(text, str):
        return text or ""
    if any(c in text for c in ("\u201c", "\u201d", "\u2018", "\u2019", "\u2014")):
        return text
    s = text
    out: list[str] = []
    open_dq = True
    for ch in s:
        if ch == '"':
            out.append("\u201c" if open_dq else "\u201d")
            open_dq = not open_dq
        else:
            out.append(ch)
    s = "".join(out)
    s = re.sub(r"(\w)'(\w)", "\\1\u2019\\2", s)
    s = re.sub(r"(?<!\w)'(\S[^']*?\S)'(?!\w)", "\u2018\\1\u2019", s)
    s = re.sub(r"(?<=\w) - (?=\w)", " \u2014 ", s)
    return s


def _slugify(text: str, *, max_len: int = 72) -> str:
    """ASCII slug for filenames. Mirrors openwebui_artifacts._slug_filename."""
    t = unicodedata.normalize("NFKD", text or "")
    t = "".join(c for c in t if not unicodedata.combining(c))
    t = t.lower()
    t = re.sub(r"[^a-z0-9]+", "-", t)
    t = re.sub(r"-{2,}", "-", t).strip("-")
    return (t or "document")[:max_len]


def _human_filename(text: str, *, max_len: int = 90) -> str:
    """Readable filename base from the document title.

    Keeps the original casing, spaces and accents so downloads look like
    ``My Report — 2026.docx`` instead of ``document-my-report_20260709_ab12``.
    Only strips characters that are illegal in filenames / break URLs.
    """
    t = (text or "").strip()
    t = re.sub(r'[\\/:*?"<>|\r\n\t]+', " ", t)   # illegal on Windows/URLs
    t = t.replace("\u2044", "-")                   # fraction slash lookalike
    t = re.sub(r"\s{2,}", " ", t).strip(" .-")
    return (t or "Document")[:max_len].strip(" .-")


def _hex_to_rgb(value: str) -> tuple[int, int, int]:
    """Convert ``#RRGGBB`` / ``RRGGBB`` to (r, g, b) tuple. Falls back to black."""
    if not isinstance(value, str):
        return (0, 0, 0)
    h = value.lstrip("#").strip()
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    if len(h) != 6:
        return (0, 0, 0)
    try:
        return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    except ValueError:
        return (0, 0, 0)


def _hex_clean(value: str) -> str:
    """Return a clean uppercase 6-char hex (no #), suitable for OOXML <w:fill>."""
    r, g, b = _hex_to_rgb(value)
    return f"{r:02X}{g:02X}{b:02X}"


def _mix_hex(a: str, b: str, t: float) -> str:
    """Linearly blend two hex colours. ``t=0`` -> a, ``t=1`` -> b."""
    ra, ga, ba = _hex_to_rgb(a)
    rb, gb, bb = _hex_to_rgb(b)
    t = max(0.0, min(1.0, t))
    return "{:02X}{:02X}{:02X}".format(
        round(ra + (rb - ra) * t),
        round(ga + (gb - ga) * t),
        round(ba + (bb - ba) * t),
    )


def _darken(color: str, t: float = 0.30) -> str:
    return _mix_hex(color, "000000", t)


def _lighten(color: str, t: float = 0.85) -> str:
    return _mix_hex(color, "FFFFFF", t)


def _relative_luminance(color: str) -> float:
    """Perceptual luminance (0..1) used to pick readable on-fill text."""
    r, g, b = (c / 255.0 for c in _hex_to_rgb(color))
    def _lin(c: float) -> float:
        return c / 12.92 if c <= 0.03928 else ((c + 0.055) / 1.055) ** 2.4
    return 0.2126 * _lin(r) + 0.7152 * _lin(g) + 0.0722 * _lin(b)


def _on_color(fill: str) -> str:
    """Return black or white — whichever reads better on ``fill``."""
    return "1A1A1A" if _relative_luminance(fill) > 0.45 else "FFFFFF"


# ── Design tokens ─────────────────────────────────────────────────────────────
# A single accent colour seeds a full, coherent palette (accent shades, rules,
# zebra fills, muted text). Every renderer reads colours from this dict via
# ``_theme(doc)`` instead of hard-coding hex values, so a document's look is
# controlled in ONE place and stays consistent end to end.
DEFAULT_ACCENT = "1E2761"

# Neutral greys (kept fixed — they read well on white regardless of accent).
_INK = "1A1A1A"        # primary body text
_INK_SOFT = "4A4E57"   # secondary text (captions, sub-labels)
_INK_MUTED = "7A7F8A"  # tertiary text (metadata, footers)


def _resolve_theme(spec: dict) -> dict:
    """Derive the full design-token palette from the resolved spec.

    Reads ``styles.accent`` (+ optional ``heading_color``/``font``/``size_pt``)
    and expands it into a coherent set of colours + type/space scales. The
    result is cached on the Document via ``_theme(doc)``.
    """
    styles = spec.get("styles") or {}
    accent = _hex_clean(styles.get("accent") or DEFAULT_ACCENT)
    heading = _hex_clean(styles.get("heading_color") or accent)
    return {
        # Brand
        "accent": accent,
        "accent_dark": _darken(accent, 0.32),
        "accent_soft": _lighten(accent, 0.88),   # faint tint for bands/panels
        "accent_tint": _lighten(accent, 0.80),   # slightly stronger tint
        "on_accent": _on_color(accent),          # text over a full accent fill
        # Text
        "heading": heading,
        "ink": _INK,
        "ink_soft": _INK_SOFT,
        "ink_muted": _INK_MUTED,
        # Structure
        "rule": _lighten(accent, 0.62),          # heading dividers, hairlines
        "rule_strong": _lighten(accent, 0.35),
        "zebra": "F5F7FA",                        # table row striping
        "cell_border": "DCE0E6",
        "panel": "F4F6F9",                        # callout/code neutral panel
        # Typography
        "font": styles.get("font") or DEFAULT_FONT,
        "mono_font": styles.get("mono_font") or "Consolas",
        "size_pt": int(styles.get("size_pt") or DEFAULT_FONT_SIZE_PT),
    }


_DEFAULT_THEME = _resolve_theme({})


def _theme(doc: _DocxDocument) -> dict:
    """Return the design tokens attached to this document (or defaults)."""
    return getattr(doc, "_neura_theme", None) or _DEFAULT_THEME


def _rgb(hex_color: str) -> "RGBColor":
    """Shorthand: hex string -> python-docx RGBColor."""
    return RGBColor(*_hex_to_rgb(hex_color))


# Inline markdown patterns: **bold**, *italic*, ~accent~, `code`, [text](url)
# Order matters: links first (the ``url`` may contain * or _), then formatting.
_INLINE_PATTERN = re.compile(
    r"(\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)]+)\))"
    r"|(\*\*(?P<bold>[^*]+?)\*\*)"
    r"|(\*(?P<italic>[^*]+?)\*)"
    r"|(~(?P<accent>[^~]+?)~)"
    r"|(`(?P<code>[^`]+?)`)"
)


def _iter_inline_runs(text: str) -> list[dict]:
    """Tokenise inline markdown into a flat list of run-spec dicts.

    Each entry: ``{"text": ..., "bold": bool, "italic": bool, "code": bool,
    "accent": bool, "link": Optional[str]}``. The Word renderer turns each
    entry into a separate ``<w:r>`` (or ``<w:hyperlink>`` for links).
    """
    if not text:
        return []
    text = _smart_quotes(str(text))
    out: list[dict] = []
    pos = 0
    for m in _INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            out.append({"text": text[pos:m.start()]})
        gd = m.groupdict()
        if gd.get("link_text") is not None:
            out.append({
                "text": gd["link_text"],
                "link": gd.get("link_url"),
            })
        elif gd.get("bold") is not None:
            out.append({"text": gd["bold"], "bold": True})
        elif gd.get("italic") is not None:
            out.append({"text": gd["italic"], "italic": True})
        elif gd.get("accent") is not None:
            out.append({"text": gd["accent"], "accent": True})
        elif gd.get("code") is not None:
            out.append({"text": gd["code"], "code": True})
        pos = m.end()
    if pos < len(text):
        out.append({"text": text[pos:]})
    return out

# endregion


# region ── Templates ─────────────────────────────────────────────────────────

# Each template is a dict of *defaults*. The renderer deep-merges the
# user-supplied document spec on top, so every template field can be
# overridden block by block. Only ``page``/``styles``/``header``/``footer``
# defaults belong here — the actual content is always provided by the model.

_TEMPLATE_BLANK: dict = {
    "page": {"size": "a4", "orientation": "portrait", "margin_mm": 25.0},
    "styles": {
        "font": DEFAULT_FONT,
        "size_pt": DEFAULT_FONT_SIZE_PT,
        "accent": "#1E2761",
        "heading_color": "#1F2430",
    },
    "header": {},
    "footer": {},
    "cover": None,
}

_TEMPLATE_LETTER: dict = {
    "page": {"size": "a4", "orientation": "portrait", "margin_mm": 30.0},
    "styles": {
        "font": "Calibri",
        "size_pt": 11,
        "accent": "#1F3864",
        "heading_color": "#1F3864",
        "heading_rule": False,  # letters read cleaner without section rules
    },
    # Letters typically don't use a running header/footer; the address block
    # lives at the top of the body. Page numbers only when explicitly asked.
    "header": {},
    "footer": {"show_page_numbers": False},
    "cover": None,
}

_TEMPLATE_REPORT: dict = {
    "page": {"size": "a4", "orientation": "portrait", "margin_mm": 24.0},
    "styles": {
        "font": "Calibri",
        "size_pt": 11,
        "accent": "#2E5AAC",
        "heading_color": "#1F2A44",
        "numbered_headings": True,
        "numbered_headings_depth": 3,
        "cover_style": "band",  # bold accent panel cover (Claude-grade default)
    },
    "header": {"show_page_numbers": False},
    "footer": {
        "show_page_numbers": True,
        "page_number_format": "{current} / {total}",
    },
    "cover": "auto",  # request a cover page when title/subtitle present
}

_TEMPLATE_MEMO: dict = {
    "page": {"size": "a4", "orientation": "portrait", "margin_mm": 24.0},
    "styles": {
        "font": "Calibri",
        "size_pt": 11,
        "accent": "#334155",
        "heading_color": "#1F2937",
        "heading_rule_levels": [1],
    },
    "header": {},
    "footer": {"show_page_numbers": True},
    "cover": None,
    # Memo header block (To:/From:/Date:/Subject:) is rendered automatically
    # at the top of the body when the spec includes ``memo_fields``.
}

_TEMPLATE_PROPOSAL: dict = {
    "page": {"size": "a4", "orientation": "portrait", "margin_mm": 22.0},
    "styles": {
        "font": "Calibri",
        "size_pt": 11,
        "accent": "#1E2761",
        "heading_color": "#1E2761",
        "cover_style": "band",  # bold accent panel cover
        "numbered_headings": False,
    },
    "header": {},
    "footer": {
        "show_page_numbers": True,
        "page_number_format": "{current} / {total}",
    },
    "cover": "auto",
}

_TEMPLATE_MINUTES: dict = {
    "page": {"size": "a4", "orientation": "portrait", "margin_mm": 24.0},
    "styles": {
        "font": "Calibri",
        "size_pt": 11,
        "accent": "#2F6F4F",
        "heading_color": "#22503A",
        "heading_rule_levels": [1, 2],
    },
    "header": {},
    "footer": {"show_page_numbers": True},
    "cover": None,
}

_TEMPLATE_WHITEPAPER: dict = {
    "page": {"size": "a4", "orientation": "portrait", "margin_mm": 26.0},
    "styles": {
        "font": "Calibri",
        "size_pt": 11,
        "accent": "#0F766E",
        "heading_color": "#123B36",
        "numbered_headings": True,
        "numbered_headings_depth": 2,
        "cover_style": "rule",
    },
    "header": {"show_page_numbers": False},
    "footer": {
        "show_page_numbers": True,
        "page_number_format": "{current} / {total}",
    },
    "cover": "auto",
}

TEMPLATES: dict[str, dict] = {
    "blank": _TEMPLATE_BLANK,
    "letter": _TEMPLATE_LETTER,
    "report": _TEMPLATE_REPORT,
    "memo": _TEMPLATE_MEMO,
    "proposal": _TEMPLATE_PROPOSAL,
    "minutes": _TEMPLATE_MINUTES,
    "whitepaper": _TEMPLATE_WHITEPAPER,
}


def _deep_merge(base: dict, override: dict) -> dict:
    """Right-biased deep merge for template defaults + user spec.

    Values in ``override`` win, except when both are dicts (then recurse).
    Lists in ``override`` replace the base list outright (no concat).
    """
    out: dict = dict(base)
    for k, v in (override or {}).items():
        if (
            k in out
            and isinstance(out[k], dict)
            and isinstance(v, dict)
        ):
            out[k] = _deep_merge(out[k], v)
        else:
            out[k] = v
    return out


def _resolve_template(spec: dict) -> dict:
    """Return the merged spec (template defaults + user overrides)."""
    name = (spec.get("template") or "blank").lower().strip()
    base = TEMPLATES.get(name, _TEMPLATE_BLANK)
    merged = _deep_merge(base, spec)
    merged["template"] = name if name in TEMPLATES else "blank"

    # Normalize the styles block + accent aliases. Authors frequently
    # mis-indent YAML frontmatter (leaving ``styles:`` empty and pushing
    # ``accent_color`` to the top level) or use ``accent_color`` instead of
    # ``accent``. Coalesce all of these so the brand color is honored and
    # ``styles`` is always a dict downstream.
    styles = merged.get("styles")
    if not isinstance(styles, dict):
        styles = {}
    if not styles.get("accent"):
        alias = (
            merged.get("accent") or merged.get("accent_color")
            or merged.get("accent_hex") or styles.get("accent_color")
        )
        if alias:
            styles["accent"] = alias
    merged["styles"] = styles
    return merged

# endregion


# region ── Markdown parser (dual-input pipeline) ─────────────────────────────
# Goal: accept Markdown-with-frontmatter as an alternative to JSON, and
# emit the EXACT same dict shape that JSON arrives in. The downstream
# pipeline (template merge, OOXML rendering, image pipeline, save) is
# completely unaware of which input format was used.
#
# Design choices:
#   - markdown-it-py (CommonMark + table) as the parsing core: it's the
#     reference implementation used by Pandoc/MkDocs/Quarto, and the model
#     produces it natively (zero prompt engineering).
#   - YAML frontmatter for ALL document-level settings (template, page,
#     styles, header/footer, cover, letter_fields, memo_fields, ...): the
#     Jekyll/Hugo/Quarto convention is universally known.
#   - Fenced divs ``::: name attrs`` for our custom blocks (callout,
#     signature, page-break, toc): Pandoc's standard, supported by
#     mdit-py-plugins out of the box.
#   - GitHub admonition syntax ``> [!note] Title`` as a second alias for
#     callouts (very common in modern markdown corpora).
#   - Inline ``==text==`` for the accent highlight (CriticMarkup), mapped
#     to the renderer's existing ``~text~`` accent runs.
#   - Image src prefixes (``unsplash:`` / ``gen:`` / ``data:`` / http(s))
#     to keep the full image pipeline available without inventing extra
#     directives.
#
# Lenient by design: the parser NEVER raises. Every failure mode (broken
# YAML frontmatter, unknown directive, malformed table, unsupported image
# prefix) degrades gracefully so the document is still produced. This is
# the explicit "evita errori" contract from the user requirements.


_MD_SINGLETON = None  # cached MarkdownIt instance


def _get_md_parser():
    """Lazily build (and cache) the MarkdownIt parser singleton.

    Returns ``None`` if the markdown-it-py / mdit-py-plugins stack is not
    installed (the JSON path keeps working in that case).
    """
    global _MD_SINGLETON
    if _MD_SINGLETON is not None:
        return _MD_SINGLETON
    if not _HAS_MD_PARSER:
        return None
    md = (
        MarkdownIt("commonmark", {"html": False})
        .enable("table")
        .enable("strikethrough")
        .use(front_matter_plugin)
    )
    # Custom block containers (Pandoc fenced divs).
    for name in ("callout", "signature", "cover", "page-break", "toc",
                 "memo", "letter", "kpi", "columns", "definition", "figure"):
        md.use(container_plugin, name)
    _MD_SINGLETON = md
    return md


def _yaml_safe_load(text: str) -> dict:
    """Parse YAML lenient: returns {} on any failure (logged, not raised)."""
    if not text or yaml is None:
        return {}
    try:
        data = yaml.safe_load(text)
    except Exception:
        traceback.print_exc()
        return {}
    return data if isinstance(data, dict) else {}


# Inline ``==text==`` (CriticMarkup highlight) -> the renderer's accent
# inline syntax ``~text~``. We only rewrite balanced pairs so partial
# typos like ``==foo`` survive as plain text.
_CRITIC_HIGHLIGHT_RE = re.compile(r"==([^=\n]+?)==")


def _md_inline_to_renderer(text: str) -> str:
    """Normalise inline markdown for the existing run renderers.

    The renderer's ``_iter_inline_runs`` already understands ``**bold**``,
    ``*italic*``, `` `code` ``, and ``[link](url)``. We only need to map
    the highlight syntax (``==x==`` -> ``~x~``) so the accent colour kicks
    in. Everything else is passed through untouched.
    """
    return _CRITIC_HIGHLIGHT_RE.sub(r"~\1~", text)


# Attribute parser for fenced div info strings, e.g.
# ``::: callout type="info" title="Heads up"`` -> ``{"type":"info","title":"Heads up"}``.
_DIV_ATTR_RE = re.compile(r'(\w+)\s*=\s*"([^"]*)"')
_DIV_ATTR_BARE_RE = re.compile(r'(\w+)\s*=\s*([^\s"]+)')


def _parse_div_attrs(info: str) -> tuple[str, dict]:
    """Split a fenced-div info string into (name, attrs).

    Examples:
        ``" callout type=\"info\" title=\"X\" "`` -> ("callout",
            {"type": "info", "title": "X"})
        ``"signature align=right"`` -> ("signature", {"align": "right"})
        ``""`` -> ("", {})
    """
    info = (info or "").strip()
    if not info:
        return "", {}
    parts = info.split(None, 1)
    name = parts[0]
    rest = parts[1] if len(parts) > 1 else ""
    attrs: dict[str, str] = {}
    for m in _DIV_ATTR_RE.finditer(rest):
        attrs[m.group(1).lower()] = m.group(2)
    # Fallback for bare key=value pairs (no quotes) that didn't match above.
    consumed = set(attrs.keys())
    for m in _DIV_ATTR_BARE_RE.finditer(rest):
        k = m.group(1).lower()
        if k not in consumed:
            attrs[k] = m.group(2)
    return name, attrs


# GitHub admonition pattern at the start of a blockquote first line.
_ADMONITION_RE = re.compile(
    r"^\s*\[!(?P<kind>note|tip|info|important|warning|caution|"
    r"warn|success|danger|error)\](?:\s+(?P<title>[^\n]+))?\s*$",
    re.IGNORECASE,
)

# Map admonition kind -> our callout `kind` palette.
_ADMONITION_KIND_MAP = {
    "note": "info",
    "tip": "success",
    "info": "info",
    "important": "warning",
    "warning": "warning",
    "warn": "warning",
    "caution": "danger",
    "danger": "danger",
    "error": "danger",
    "success": "success",
}


def _image_block_from_md(alt: str, src: str, title: str) -> dict:
    """Map an inline ``![alt](src "title")`` to an internal image block.

    Recognised src prefixes (case-insensitive):
        - ``unsplash:<query>`` -> ``image_hint``
        - ``gen:<prompt>`` -> ``image_generate``
        - ``data:image/...;base64,...`` -> ``image_b64``
        - ``http(s)://...`` -> ``image_url``
    Anything else falls through as ``image_url`` if it looks like a URL,
    otherwise the block is annotated with ``_unresolved`` so the caller
    can render a placeholder paragraph.
    """
    block: dict = {"type": "image"}
    if alt:
        block["alt"] = alt
    if title:
        block["caption"] = title
    src_low = (src or "").strip()
    if not src_low:
        block["_unresolved"] = "missing src"
        return block
    low = src_low.lower()
    if low.startswith("unsplash:"):
        block["image_hint"] = src_low.split(":", 1)[1].strip()
    elif low.startswith("gen:"):
        block["image_generate"] = src_low.split(":", 1)[1].strip()
    elif low.startswith("data:image/"):
        block["image_b64"] = src_low
    elif low.startswith(("http://", "https://")):
        block["image_url"] = src_low
    else:
        block["_unresolved"] = src_low
    return block


# Frontmatter keys that map directly to the spec dict.
_FRONTMATTER_KEYS = {
    "title", "subtitle", "author", "date", "template",
    "page", "styles", "header", "footer", "cover",
    "letter_fields", "memo_fields",
    # Accent-color aliases: authors frequently mis-indent frontmatter
    # (leaving ``styles:`` empty and pushing ``accent_color`` to the top
    # level). Keep these so `_resolve_template` can fold them into
    # ``styles.accent`` instead of silently dropping the brand color.
    "accent", "accent_color", "accent_hex",
}


def _walk_inline_only_image(content: str) -> Optional[tuple[str, str, str]]:
    """If a paragraph's inline content is exactly one image, extract it.

    Returns ``(alt, src, title)`` or ``None``. We use a strict regex so
    paragraphs with mixed text + image are rendered as paragraphs (and
    the inline renderer takes care of any embedded image syntax via the
    standard inline parser).
    """
    # Note: src may contain spaces (e.g. ``unsplash:business team``), so we
    # accept anything except ``)`` and the optional ``"title"`` suffix.
    m = re.fullmatch(
        r'!\[(?P<alt>[^\]]*)\]'
        r'\((?P<src>[^\s)][^\)]*?)(?:\s+"(?P<title>[^"]*)")?\)',
        content.strip(),
    )
    if not m:
        return None
    return m.group("alt") or "", m.group("src"), m.group("title") or ""


def _parse_table_tokens(tokens: list, start: int) -> tuple[dict, int]:
    """Walk a markdown-it table_open .. table_close range into a table block.

    Returns ``({"type":"table","headers":[...],"rows":[[...]],...}, end_idx)``.
    """
    headers: list[str] = []
    rows: list[list[str]] = []
    align: list[str] = []
    current_row: list[str] = []
    in_header = False
    i = start + 1
    while i < len(tokens):
        t = tokens[i]
        if t.type == "table_close":
            return (
                {
                    "type": "table",
                    "headers": headers,
                    "rows": rows,
                    "shaded_header": True,
                    "zebra": True,
                    **({"align": align} if any(a for a in align) else {}),
                },
                i,
            )
        if t.type == "thead_open":
            in_header = True
        elif t.type == "thead_close":
            in_header = False
        elif t.type == "tr_open":
            current_row = []
        elif t.type in ("th_open", "td_open"):
            if in_header:
                style = (t.attrs or {}).get("style", "")
                m = re.search(r"text-align\s*:\s*(left|center|right)", style)
                align.append(m.group(1) if m else "")
        elif t.type == "inline":
            current_row.append(_md_inline_to_renderer(t.content or ""))
        elif t.type == "tr_close":
            if in_header:
                headers = current_row
            else:
                rows.append(current_row)
        i += 1
    # Auto-closed by EOF — return what we have.
    return (
        {
            "type": "table",
            "headers": headers,
            "rows": rows,
            "shaded_header": True,
            "zebra": True,
        },
        len(tokens) - 1,
    )


def _parse_list_items(tokens: list, start: int, *, ordered: bool) -> tuple[dict, int]:
    """Walk a list_open .. list_close range into our list block.

    Returns ``({"type":"list","ordered":bool,"items":[...]}, end_idx)``. List
    items support nested lists (recursive) and inline-formatted text.
    """
    close_type = "ordered_list_close" if ordered else "bullet_list_close"
    items: list = []
    i = start + 1
    while i < len(tokens):
        t = tokens[i]
        if t.type == close_type:
            return (_finalize_list_block(ordered, items), i)
        if t.type == "list_item_open":
            text_parts: list[str] = []
            sub_items: list = []
            j = i + 1
            while j < len(tokens) and tokens[j].type != "list_item_close":
                tt = tokens[j]
                if tt.type == "inline":
                    text_parts.append(_md_inline_to_renderer(tt.content or ""))
                elif tt.type in ("bullet_list_open", "ordered_list_open"):
                    sub_block, end = _parse_list_items(
                        tokens, j, ordered=(tt.type == "ordered_list_open"),
                    )
                    sub_items.append(sub_block)
                    j = end
                j += 1
            text = " ".join(p for p in text_parts if p)
            if sub_items:
                items.append({
                    "text": text,
                    "items": [
                        si.get("items", []) if isinstance(si, dict) else si
                        for si in sub_items
                    ][0] if len(sub_items) == 1 else [
                        si for sb in sub_items for si in sb.get("items", [])
                    ],
                })
            else:
                items.append(text)
            i = j
        i += 1
    return (_finalize_list_block(ordered, items), len(tokens) - 1)


_TASK_ITEM_RE = re.compile(r"^\[(?P<mark>[ xX\u2713])\]\s+(?P<rest>.*)$", re.DOTALL)


def _finalize_list_block(ordered: bool, items: list) -> dict:
    """Emit a list block, upgrading GitHub task-lists (``- [ ]``) to checklist.

    If EVERY top-level string item begins with a ``[ ]``/``[x]`` marker the
    list becomes a checklist block (rendered with real ☐/☑ glyphs).
    """
    str_items = [it for it in items if isinstance(it, str)]
    if (
        not ordered and items and len(str_items) == len(items)
        and all(_TASK_ITEM_RE.match(it.strip()) for it in str_items)
    ):
        checklist: list = []
        for it in str_items:
            m = _TASK_ITEM_RE.match(it.strip())
            checked = m.group("mark").lower() in ("x", "\u2713")
            checklist.append({"text": m.group("rest").strip(), "checked": checked})
        return {"type": "list", "list_style": "checklist", "items": checklist}
    return {"type": "list", "ordered": ordered, "items": items}


def _parse_blockquote(tokens: list, start: int) -> tuple[dict, int]:
    """Walk a blockquote_open .. blockquote_close range.

    If the first line matches the GitHub admonition pattern
    ``[!kind] Title``, emit a callout block; otherwise emit a quote block.
    """
    paragraphs: list[str] = []
    i = start + 1
    while i < len(tokens):
        t = tokens[i]
        if t.type == "blockquote_close":
            break
        if t.type == "inline":
            paragraphs.append(t.content or "")
        i += 1
    full = "\n".join(paragraphs).strip()
    first_line, _, rest = full.partition("\n")
    m = _ADMONITION_RE.match(first_line)
    if m:
        kind = _ADMONITION_KIND_MAP.get(m.group("kind").lower(), "info")
        title = (m.group("title") or "").strip()
        body = _md_inline_to_renderer(rest.strip())
        block: dict = {"type": "callout", "kind": kind, "text": body}
        if title:
            block["title"] = title
        return block, i
    return {"type": "quote", "text": _md_inline_to_renderer(full)}, i


def _container_raw_body(open_token, src_lines: Optional[list]) -> str:
    """Extract the verbatim body lines of a fenced div using its ``.map``.

    markdown-it flattens inner content into inline tokens, which loses the
    original YAML/structured layout. For structured containers (kpi, columns,
    definition) we need the raw text, so we slice the source by line map.
    """
    if not src_lines or not getattr(open_token, "map", None):
        return ""
    start_line, end_line = open_token.map[0], open_token.map[1]
    # Body is between the opening ``:::`` line and the closing ``:::`` line.
    body_lines = src_lines[start_line + 1:end_line]
    while body_lines and body_lines[-1].strip().startswith(":::"):
        body_lines = body_lines[:-1]
    return "\n".join(body_lines).strip()


def _parse_container(
    tokens: list, start: int, name: str, src_lines: Optional[list] = None,
) -> tuple[Optional[dict], int]:
    """Walk a container_<name>_open .. container_<name>_close range.

    Returns the corresponding block dict (or None to skip), plus end_idx.
    """
    close_type = f"container_{name}_close"
    open_token = tokens[start]
    _, attrs = _parse_div_attrs(open_token.info or "")
    inner_inlines: list[str] = []
    i = start + 1
    while i < len(tokens):
        t = tokens[i]
        if t.type == close_type:
            break
        if t.type == "inline":
            inner_inlines.append(t.content or "")
        i += 1
    body_raw = "\n\n".join(p for p in inner_inlines if p.strip())
    body = _md_inline_to_renderer(body_raw)

    # Structured containers: parse the verbatim YAML body.
    if name in ("kpi", "columns", "definition", "figure"):
        raw = _container_raw_body(open_token, src_lines) or body_raw
        data = _yaml_safe_load(raw)
        if name == "kpi":
            items = data.get("items") if isinstance(data, dict) else data
            return {"type": "kpi", "items": items or []}, i
        if name == "columns":
            cols = data.get("columns") if isinstance(data, dict) else data
            return {"type": "columns", "columns": cols or []}, i
        if name == "definition":
            items = data.get("items") if isinstance(data, dict) else data
            if not items:
                # Fallback: "Term: description" per line.
                items = []
                for line in raw.splitlines():
                    if ":" in line:
                        term, _, desc = line.partition(":")
                        items.append({"term": term.strip(), "description": desc.strip()})
            return {"type": "definition", "items": items or []}, i
        if name == "figure":
            blk: dict = {"type": "figure"}
            if isinstance(data, dict):
                for k in ("src", "image", "image_b64", "image_url", "image_hint",
                          "image_generate", "caption", "alt", "width_mm"):
                    if data.get(k) is not None:
                        blk[k if k not in ("src", "image") else "_src"] = data[k]
                if "_src" in blk:
                    src = str(blk.pop("_src"))
                    blk.update(_image_block_from_md(blk.get("alt", ""), src, ""))
                    blk["type"] = "figure"
            if attrs.get("caption") and not blk.get("caption"):
                blk["caption"] = attrs["caption"]
            return blk, i

    if name == "callout":
        kind = attrs.get("type") or attrs.get("kind") or "info"
        block: dict = {"type": "callout", "kind": kind.lower(), "text": body}
        if attrs.get("title"):
            block["title"] = attrs["title"]
        return block, i
    if name == "signature":
        # A signature accepts its fields either as fence attributes
        # (``::: signature name="..." role="..." date="..."``) or as body
        # lines (name / role / date, one per line). Attributes win; body
        # lines fill any field the attributes didn't set.
        lines = [ln.strip() for ln in body_raw.splitlines() if ln.strip()]
        block = {"type": "signature"}
        if attrs.get("name"):
            block["name"] = attrs["name"]
        elif lines:
            block["name"] = lines[0]
        if attrs.get("role"):
            block["role"] = attrs["role"]
        elif len(lines) >= 2:
            block["role"] = lines[1]
        if attrs.get("date"):
            block["date"] = attrs["date"]
        elif len(lines) >= 3:
            block["date"] = lines[2]
        if attrs.get("align"):
            block["align"] = attrs["align"]
        return block, i
    if name == "page-break":
        return {"type": "page_break"}, i
    if name == "toc":
        block = {"type": "toc"}
        if attrs.get("title"):
            block["title"] = attrs["title"]
        if attrs.get("depth"):
            try:
                block["depth"] = int(attrs["depth"])
            except ValueError:
                pass
        return block, i
    if name in ("memo", "letter", "cover"):
        # These are spec-level overrides expressed as containers; the
        # parser surfaces them as side-channel blocks the document builder
        # consumes from the spec dict (memo_fields/letter_fields/cover).
        # We return a sentinel block carrying the parsed YAML body.
        spec_dict = _yaml_safe_load(body_raw)
        return {"_meta": name, "data": spec_dict}, i
    # Unknown container: degrade gracefully to a quote (visible but
    # marked) so the document still renders.
    return {
        "type": "quote",
        "text": f"[{name}] {body}",
    }, i


def _parse_markdown(text: str) -> dict:
    """Parse Markdown-with-frontmatter into the internal spec dict.

    Lenient: never raises. Failures degrade to best-effort output (the
    document is always produced).
    """
    md = _get_md_parser()
    if md is None:
        # No markdown stack available — return a minimal spec carrying the
        # raw text as a single paragraph so the user still gets a docx.
        return {
            "blocks": [{"type": "paragraph", "text": text or ""}],
        }

    try:
        tokens = md.parse(text or "")
    except Exception:
        traceback.print_exc()
        return {"blocks": [{"type": "paragraph", "text": text or ""}]}

    src_lines = (text or "").split("\n")
    spec: dict = {}
    blocks: list = []

    # 1. Pull frontmatter if present (always the first token if any).
    if tokens and tokens[0].type == "front_matter":
        fm_data = _yaml_safe_load(tokens[0].content or "")
        for k, v in fm_data.items():
            if k in _FRONTMATTER_KEYS:
                spec[k] = v
        # Unknown frontmatter keys are silently ignored (lenient).

    # 2. Walk the rest of the token stream at depth 0 only.
    i = 0
    while i < len(tokens):
        t = tokens[i]
        if t.level != 0:
            i += 1
            continue
        ttype = t.type
        try:
            if ttype == "front_matter":
                pass

            elif ttype == "heading_open":
                level = int(t.tag[1]) if t.tag.startswith("h") else 1
                inline = tokens[i + 1].content if i + 1 < len(tokens) else ""
                blocks.append({
                    "type": "heading",
                    "level": min(max(level, 1), 4),
                    "text": _md_inline_to_renderer(inline),
                })
                i += 2  # skip inline + heading_close handled by loop

            elif ttype == "paragraph_open":
                inline = tokens[i + 1].content if i + 1 < len(tokens) else ""
                stripped = inline.strip()
                # Special single-line paragraphs we treat as directives.
                if stripped == "[[toc]]":
                    blocks.append({"type": "toc"})
                elif stripped in ("\\newpage", "\\\\newpage"):
                    blocks.append({"type": "page_break"})
                elif re.fullmatch(r":{3,}", stripped):
                    # Dangling fenced-div marker (unmatched ``:::``) — LLMs emit
                    # these often; swallow it instead of printing literal colons.
                    pass
                else:
                    img = _walk_inline_only_image(stripped)
                    if img is not None:
                        alt, src, title = img
                        block = _image_block_from_md(alt, src, title)
                        if "_unresolved" in block:
                            blocks.append({
                                "type": "paragraph",
                                "text": f"[image unavailable: "
                                        f"{block['_unresolved']}]",
                            })
                        else:
                            blocks.append(block)
                    else:
                        blocks.append({
                            "type": "paragraph",
                            "text": _md_inline_to_renderer(inline),
                        })
                i += 2  # inline + paragraph_close

            elif ttype == "bullet_list_open":
                blk, end = _parse_list_items(tokens, i, ordered=False)
                blocks.append(blk)
                i = end

            elif ttype == "ordered_list_open":
                blk, end = _parse_list_items(tokens, i, ordered=True)
                blocks.append(blk)
                i = end

            elif ttype == "table_open":
                blk, end = _parse_table_tokens(tokens, i)
                blocks.append(blk)
                i = end

            elif ttype == "blockquote_open":
                blk, end = _parse_blockquote(tokens, i)
                blocks.append(blk)
                i = end

            elif ttype == "hr":
                blocks.append({"type": "horizontal_rule"})

            elif ttype in ("fence", "code_block"):
                info_parts = (t.info or "").strip().split(None, 1) if ttype == "fence" else []
                blocks.append({
                    "type": "code",
                    "code": t.content or "",
                    "language": info_parts[0] if info_parts else "",
                })

            elif ttype.startswith("container_") and ttype.endswith("_open"):
                name = ttype[len("container_"):-len("_open")]
                blk, end = _parse_container(tokens, i, name, src_lines)
                if blk is not None:
                    if blk.get("_meta") in ("memo", "letter", "cover"):
                        # Side-channel: merge into the spec dict, NOT into
                        # blocks. ``cover`` becomes spec["cover"]; memo /
                        # letter become memo_fields / letter_fields.
                        meta = blk["_meta"]
                        data = blk.get("data") or {}
                        if meta == "cover":
                            spec["cover"] = data or "auto"
                        elif meta == "memo":
                            spec["memo_fields"] = data
                        elif meta == "letter":
                            spec["letter_fields"] = data
                    else:
                        blocks.append(blk)
                i = end

            # Skip stray close tokens / fence tokens we don't render.
        except Exception:
            traceback.print_exc()
            # On any unexpected parser fault, append a placeholder so the
            # user notices something off but the document still builds.
            blocks.append({
                "type": "paragraph",
                "text": "[unrecognized block skipped]",
            })
        i += 1

    spec["blocks"] = blocks
    return spec

# endregion


# region ── Low-level OOXML helpers ────────────────────────────────────────────
# python-docx exposes high-level wrappers for ~80% of features. The remaining
# 20% (numbering definitions, shading, custom borders, page numbers) requires
# poking the OOXML element tree directly. These helpers keep that ugliness
# isolated — every renderer above just calls them.

def _set_cell_shading(cell, hex_fill: str) -> None:
    """Apply a solid fill colour to a table cell.

    Anthropic skill rule: use ``ShadingType.CLEAR`` (i.e. <w:shd val="clear">)
    rather than SOLID — SOLID renders as opaque black on Google Docs and some
    older Word versions because it interprets ``fill`` as the foreground.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex_clean(hex_fill))
    # Replace any existing shading
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    tc_pr.append(shd)


def _set_cell_borders(cell, *, color: str = "CCCCCC", size_pt: int = 4) -> None:
    """Set thin uniform borders on a single cell. ``size`` is in 1/8 pt."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    color = _hex_clean(color)
    for side in ("top", "left", "bottom", "right"):
        b = tc_borders.find(qn(f"w:{side}"))
        if b is None:
            b = OxmlElement(f"w:{side}")
            tc_borders.append(b)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size_pt))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)


def _set_cell_margins(
    cell,
    *,
    top: int = 80,
    bottom: int = 80,
    left: int = 120,
    right: int = 120,
) -> None:
    """Set internal cell margins in DXA (1/20 pt). Default ~4mm/6mm."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def _set_paragraph_bottom_border(
    paragraph, *, color: str = "2E75B6", size_pt: int = 6
) -> None:
    """Add a thin bottom border to a paragraph (use as a divider rule).

    Anthropic skill: never use 1-row tables as dividers — they have a min
    height and render as visible empty boxes. Use a paragraph border instead.
    """
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        p_pr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pbdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _hex_clean(color))


def _add_field(
    paragraph,
    instr_text: str,
    *,
    placeholder: str = "1",
    dirty: bool = False,
) -> None:
    """Insert a Word field code (e.g. PAGE, NUMPAGES, TOC) into a paragraph.

    Word fields are 4-element runs: fldChar(begin) + instrText + fldChar(separate)
    + fldChar(end). python-docx has no high-level helper for them.

    Set ``dirty=True`` so Word/LibreOffice will recompute the field result on
    next open even if the document doesn't have the global ``updateFields``
    setting enabled — this is what makes the TOC actually populate without
    asking the user to right-click → Update Field.
    """
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    if dirty:
        fld_begin.set(qn("w:dirty"), "true")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    # Result placeholder shown until the consumer first updates fields. For
    # PAGE/NUMPAGES "1" is fine; for TOC we show a hint instead of a single
    # digit so the document never looks broken if a viewer can't resolve
    # fields (e.g. preview tools that don't recompute).
    result = OxmlElement("w:t")
    result.text = placeholder
    run._r.append(result)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _set_update_fields_on_open(doc: _DocxDocument) -> None:
    """Mark ``settings.xml`` so Word recomputes all fields when the doc opens.

    Without this, TOC fields (and PAGE/NUMPAGES too in some viewers) stay on
    their literal placeholder value because Word treats fields as cached
    results. The TOC field by itself just declares the intent — the actual
    list of headings is generated when the field is updated. Setting
    ``<w:updateFields w:val="true"/>`` triggers that update automatically the
    first time the user opens the document.

    Idempotent: running multiple times only sets the flag once.
    """
    try:
        settings_el = doc.settings.element
    except Exception:
        return
    existing = settings_el.find(qn("w:updateFields"))
    if existing is not None:
        existing.set(qn("w:val"), "true")
        return
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    # Schema is permissive about position; appending is safe.
    settings_el.append(upd)


def _add_toc_field(paragraph, *, depth: int = 3) -> None:
    """Insert a TOC field. Word repopulates it on first open ("Update TOC").

    Combined with :func:`_set_update_fields_on_open` and the per-field
    ``dirty`` flag, the TOC populates automatically without the user having
    to right-click → Update Field.
    """
    instr = (
        rf'TOC \o "1-{depth}" \h \z \u'
    )
    _add_field(
        paragraph,
        instr,
        placeholder="Right-click and select 'Update Field' to populate the table of contents.",
        dirty=True,
    )


def _ensure_numbering_definitions(doc: _DocxDocument, accent_hex: str) -> dict[str, int]:
    """Create numbering definitions for bullets + decimal lists, return numId map.

    python-docx ships with no high-level numbering API; we have to inject the
    abstractNum + num XML into ``word/numbering.xml`` (creating the part if
    missing). Returns ``{"bullet": int, "decimal": int}`` — the numIds the
    block renderer will reference via ``<w:numPr>``.

    Idempotent: running twice yields the same numIds.
    """
    # Cache on the document instance
    cached = getattr(doc, "_neura_num_ids", None)
    if cached:
        return cached

    # Get or create the numbering part
    try:
        numbering = doc.part.numbering_part.element  # type: ignore[attr-defined]
    except (AttributeError, KeyError, NotImplementedError):
        # Some python-docx versions auto-create on access via .numbering_part
        from docx.parts.numbering import NumberingPart  # type: ignore
        numbering_part = NumberingPart.new()
        doc.part.relate_to(numbering_part, doc.part.numbering_part_relationship_type if False else "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering")
        numbering = numbering_part.element

    nsmap = numbering.nsmap

    def _w(tag: str) -> str:
        return f"{{{nsmap.get('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')}}}{tag}"

    # Pick fresh abstractNumId / numId values
    existing_abstract = numbering.findall(_w("abstractNum"))
    existing_nums = numbering.findall(_w("num"))
    abstract_ids = [int(e.get(_w("abstractNumId")) or 0) for e in existing_abstract]
    num_ids = [int(e.get(_w("numId")) or 0) for e in existing_nums]
    next_abstract = (max(abstract_ids) + 1) if abstract_ids else 0
    next_num = (max(num_ids) + 1) if num_ids else 1

    # Indent scale: comfortable but not sprawling (0.5in base, 0.32in step).
    _base_ind = 460
    _step_ind = 360

    def _make_abstract(abs_id: int, kind: str) -> "OxmlElement":
        an = OxmlElement("w:abstractNum")
        an.set(qn("w:abstractNumId"), str(abs_id))
        glyphs = ["\u2022", "\u25E6", "\u25AA", "\u25AB"]  # • ◦ ▪ ▫
        for level in range(4):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(level))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)
            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
            lvl.append(num_fmt)
            lvl_text = OxmlElement("w:lvlText")
            if kind == "bullet":
                lvl_text.set(qn("w:val"), glyphs[level % len(glyphs)])
            elif kind == "legal":
                # Cumulative legal numbering: 1 / 1.1 / 1.1.1 / 1.1.1.1
                lvl_text.set(
                    qn("w:val"),
                    ".".join("%" + str(k + 1) for k in range(level + 1)),
                )
            else:  # decimal
                lvl_text.set(qn("w:val"), "%" + str(level + 1) + ".")
            lvl.append(lvl_text)
            lvl_jc = OxmlElement("w:lvlJc")
            lvl_jc.set(qn("w:val"), "left")
            lvl.append(lvl_jc)
            ppr = OxmlElement("w:pPr")
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(_base_ind + _step_ind * level))
            ind.set(qn("w:hanging"), "300" if kind != "legal" else "360")
            ppr.append(ind)
            lvl.append(ppr)
            if kind == "bullet":
                rpr = OxmlElement("w:rPr")
                rfonts = OxmlElement("w:rFonts")
                # Level 0 uses Symbol (round •); deeper levels keep the glyph.
                fnt = "Symbol" if level == 0 else "Arial"
                rfonts.set(qn("w:ascii"), fnt)
                rfonts.set(qn("w:hAnsi"), fnt)
                rpr.append(rfonts)
                lvl.append(rpr)
            an.append(lvl)
        return an

    def _make_num(num_id: int, abs_id: int) -> "OxmlElement":
        n = OxmlElement("w:num")
        n.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abs_id))
        n.append(abstract_ref)
        return n

    abs_bullet = next_abstract
    abs_decimal = next_abstract + 1
    abs_legal = next_abstract + 2
    num_bullet = next_num
    num_decimal = next_num + 1
    num_legal = next_num + 2

    # abstractNum elements must come BEFORE num elements per OOXML schema.
    numbering.insert(0, _make_abstract(abs_legal, "legal"))
    numbering.insert(0, _make_abstract(abs_decimal, "decimal"))
    numbering.insert(0, _make_abstract(abs_bullet, "bullet"))
    numbering.append(_make_num(num_bullet, abs_bullet))
    numbering.append(_make_num(num_decimal, abs_decimal))
    numbering.append(_make_num(num_legal, abs_legal))

    result = {
        "bullet": num_bullet,
        "decimal": num_decimal,
        "legal": num_legal,
        "_abstracts": {
            "bullet": abs_bullet, "decimal": abs_decimal, "legal": abs_legal,
        },
        "_next_num": num_legal + 1,
    }
    doc._neura_num_ids = result  # type: ignore[attr-defined]
    return result


def _mint_list_num(doc: _DocxDocument, kind: str) -> int:
    """Create a fresh ``<w:num>`` for an abstract so numbering RESTARTS.

    Ordered/legal lists that are separate blocks should each start at 1;
    OOXML restarts when a new numId references the same abstractNum. Bullet
    lists don't need this (order is irrelevant), so callers pass the shared id.
    """
    ids = getattr(doc, "_neura_num_ids", None)
    if not ids:
        return 1
    abstracts = ids.get("_abstracts") or {}
    abs_id = abstracts.get(kind)
    if abs_id is None:
        return ids.get(kind, 1)
    try:
        numbering = doc.part.numbering_part.element  # type: ignore[attr-defined]
    except Exception:
        return ids.get(kind, 1)
    new_num = ids["_next_num"]
    ids["_next_num"] = new_num + 1
    n = OxmlElement("w:num")
    n.set(qn("w:numId"), str(new_num))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abs_id))
    n.append(ref)
    numbering.append(n)
    return new_num

# endregion


# region ── Document setup (page, styles, header/footer, cover) ───────────────

def _apply_page_setup(doc: _DocxDocument, page: dict) -> None:
    """Apply page size + orientation + margins to all sections.

    Anthropic rule: ALWAYS set page size explicitly (default differs between
    docx-js and python-docx; python-docx defaults to Letter, we want A4).
    Landscape: swap width/height after setting orientation (python-docx
    behaves differently from docx-js here — it does NOT auto-swap).
    """
    size = (page.get("size") or DEFAULT_PAGE_SIZE).lower()
    orientation = (page.get("orientation") or "portrait").lower()
    margin_mm = float(page.get("margin_mm") or DEFAULT_MARGIN_MM)
    w_mm, h_mm = PAGE_SIZES_MM.get(size, PAGE_SIZES_MM[DEFAULT_PAGE_SIZE])

    for section in doc.sections:
        if orientation == "landscape":
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = Mm(h_mm)
            section.page_height = Mm(w_mm)
        else:
            section.orientation = WD_ORIENTATION.PORTRAIT
            section.page_width = Mm(w_mm)
            section.page_height = Mm(h_mm)
        section.top_margin = Mm(margin_mm)
        section.bottom_margin = Mm(margin_mm)
        section.left_margin = Mm(margin_mm)
        section.right_margin = Mm(margin_mm)
        # Header/footer distance: half the body margin keeps proportions tidy
        section.header_distance = Mm(max(10.0, margin_mm * 0.5))
        section.footer_distance = Mm(max(10.0, margin_mm * 0.5))


def _ppr_flag(ppr, tag: str, val: str | None = "true") -> None:
    """Idempotently set a boolean/simple flag element inside a pPr/rPr."""
    node = ppr.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        ppr.append(node)
    if val is not None:
        node.set(qn("w:val"), val)


def _ppr_bottom_border(ppr, *, color: str, size_pt: int = 6, space: int = 6) -> None:
    """Attach a bottom hairline to a *style* pPr (section-divider rule).

    Inserts ``<w:pBdr>`` in schema order (before shd/tabs/spacing/ind/jc/rPr)
    so strict Word does not reject or silently repair the style.
    """
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        anchor = None
        for tag in ("w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr"):
            node = ppr.find(qn(tag))
            if node is not None:
                anchor = node
                break
        if anchor is not None:
            anchor.addprevious(pbdr)
        else:
            ppr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pbdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), _hex_clean(color))


def _set_run_tracking(run, twentieths: int) -> None:
    """Add letter-spacing (tracking) to a run. ``twentieths`` = 1/20 pt."""
    try:
        rpr = run._r.get_or_add_rPr()
        spacing = rpr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            rpr.append(spacing)
        spacing.set(qn("w:val"), str(int(twentieths)))
    except Exception:
        pass


def _apply_styles(doc: _DocxDocument, styles: dict) -> None:
    """Configure base fonts + heading sizes + colours from the design tokens.

    Sets:
      - Document default font (Normal style) + widow/orphan control
      - Heading 1..4 sizes/weights/colours, with an automatic hairline rule
        beneath section-level headings (H1/H2 by default)
      - A ``Caption`` and refined ``Quote`` look
      - Hyperlink style (so links look like links)
    """
    theme = _theme(doc)
    font_name = styles.get("font") or theme["font"]
    base_size = int(styles.get("size_pt") or theme["size_pt"])
    heading_hex = theme["heading"]
    rule_hex = theme["rule_strong"]

    # Which section levels get the divider rule (overridable via styles).
    rule_levels = styles.get("heading_rule_levels")
    if rule_levels is None:
        rule_levels = list(HEADING_RULE_LEVELS) if styles.get(
            "heading_rule", True
        ) else []
    rule_levels = {int(x) for x in rule_levels}

    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(base_size)
    normal.font.color.rgb = _rgb(theme["ink"])
    npf = normal.paragraph_format
    npf.line_spacing = DEFAULT_LINE_SPACING
    npf.space_after = Pt(6)
    # East Asian + Complex Script font hints — needed for true cross-platform
    # rendering in mixed-script documents
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    # Widow/orphan control on the base style (inherited everywhere).
    _ppr_flag(normal.element.get_or_add_pPr(), "w:widowControl", "true")

    # Headings 1..4
    for level in (1, 2, 3, 4):
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        style.font.name = font_name
        style.font.size = Pt(HEADING_SIZES_PT.get(level, base_size))
        style.font.bold = True
        # H1 anchors the section in the heading colour (navy); sub-headings
        # (H2+) use the accent colour so they read as a distinct, lighter tier.
        style.font.color.rgb = _rgb(heading_hex if level == 1 else theme["accent"])
        # Spacing before/after — done via paragraph format
        pf = style.paragraph_format
        pf.space_before = Pt(HEADING_BEFORE_PT.get(level, 6))
        pf.space_after = Pt(HEADING_AFTER_PT.get(level, 4))
        pf.keep_with_next = True
        ppr = style.element.get_or_add_pPr()
        # Keep the whole heading + its first body lines together.
        _ppr_flag(ppr, "w:keepNext", "true")
        _ppr_flag(ppr, "w:keepLines", "true")
        # outlineLevel is required for TOC; python-docx doesn't expose it
        # directly so we set it via XML
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            ppr.append(outline)
        outline.set(qn("w:val"), str(level - 1))
        # Section-divider hairline beneath H1/H2 (the "polished report" look).
        if level in rule_levels:
            _ppr_bottom_border(
                ppr, color=rule_hex, size_pt=6 if level == 1 else 4, space=6,
            )

    # Caption style (figure/table captions)
    _ensure_caption_style(doc, font_name)

    # Hyperlink character style (Word's built-in)
    try:
        link_style = doc.styles["Hyperlink"]
        link_style.font.color.rgb = _rgb(theme["accent"])
        link_style.font.underline = True
    except KeyError:
        # Style is created on demand when first hyperlink is inserted
        pass


def _ensure_caption_style(doc: _DocxDocument, font_name: str) -> None:
    """Create/refine a ``Caption`` paragraph style for figures/tables."""
    theme = _theme(doc)
    try:
        style = doc.styles["Caption"]
    except KeyError:
        try:
            style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        except Exception:
            return
    style.font.name = font_name
    style.font.size = Pt(9)
    style.font.italic = True
    style.font.bold = False
    style.font.color.rgb = _rgb(theme["ink_soft"])
    pf = style.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(10)
    pf.keep_with_next = False


def _dual_tab_stops(paragraph, width_dxa: int) -> None:
    """Add centre + right tab stops for a 3-zone header/footer line."""
    ts = paragraph.paragraph_format.tab_stops
    ts.add_tab_stop(
        Emu(int((width_dxa / 2) * 914400 / 1440)), WD_TAB_ALIGNMENT.CENTER,
    )
    ts.add_tab_stop(
        Emu(int(width_dxa * 914400 / 1440)), WD_TAB_ALIGNMENT.RIGHT,
    )


# Page-number placeholders we understand, in every common spelling. Doubled
# braces are matched before single ones so ``{{page}}`` isn't split wrongly.
_PAGE_CUR_TOKENS = {"{current}", "{page}", "{{page}}", "{n}", "{{n}}"}
_PAGE_TOT_TOKENS = {"{total}", "{pages}", "{{pages}}", "{{total}}", "{N}", "{{N}}"}
_PAGENUM_SPLIT_RE = re.compile(
    r"(\{\{\s*pages?\s*\}\}|\{\{\s*total\s*\}\}|\{\{\s*[nN]\s*\}\}|"
    r"\{\s*current\s*\}|\{\s*total\s*\}|\{\s*pages?\s*\}|\{\s*[nN]\s*\})"
)


def _has_pagenum_placeholder(text: str) -> bool:
    return bool(text) and bool(_PAGENUM_SPLIT_RE.search(text))


def _footer_add_pagenum_runs(para, fmt: str, color_hex: str) -> None:
    """Expand page-number placeholders into live PAGE / NUMPAGES fields.

    Accepts ``{current}``/``{total}`` as well as ``{page}``/``{pages}`` and
    their double-brace variants (``{{page}}``/``{{pages}}``), so footers
    authored in different conventions all render real page numbers.
    """
    for part in _PAGENUM_SPLIT_RE.split(fmt):
        if not part:
            continue
        token = part.replace(" ", "")
        if token in _PAGE_CUR_TOKENS:
            _add_field(para, "PAGE")
        elif token in _PAGE_TOT_TOKENS:
            _add_field(para, "NUMPAGES")
        else:
            run = para.add_run(part)
            run.font.size = Pt(9)
            run.font.color.rgb = _rgb(color_hex)


def _build_header_footer(
    doc: _DocxDocument,
    *,
    header_spec: Optional[dict],
    footer_spec: Optional[dict],
    accent_hex: str,
    header_logo: Optional[bytes] = None,
    first_page_different: bool = False,
) -> None:
    """Populate header + footer for the primary section.

    header_spec: ``{text, logo (b64/url/hint), show_page_numbers, alignment}``
    footer_spec: ``{text | left/center/right, show_page_numbers,
                    page_number_format, alignment}``

    When ``first_page_different`` is set (e.g. a cover page is present) the
    running header/footer is suppressed on page 1 so the cover stays clean.
    """
    theme = _theme(doc)
    section = doc.sections[0]
    width_dxa = _content_width_dxa(doc) or 9360
    muted = theme["ink_muted"]

    if first_page_different:
        section.different_first_page_header_footer = True

    if header_spec:
        header = section.header
        header.is_linked_to_previous = False
        para = header.paragraphs[0]
        para.text = ""
        align = (header_spec.get("alignment") or "left").lower()
        para.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        text = (header_spec.get("text") or "").strip()
        h_left = _smart_quotes((header_spec.get("left") or "").strip())
        h_center = _smart_quotes((header_spec.get("center") or "").strip())
        h_right = _smart_quotes((header_spec.get("right") or "").strip())
        has_pnum = bool(header_spec.get("show_page_numbers"))

        def _hrun(txt: str) -> None:
            if _has_pagenum_placeholder(txt):
                _footer_add_pagenum_runs(para, txt, muted)
            else:
                r = para.add_run(txt)
                r.font.size = Pt(9)
                r.font.color.rgb = _rgb(muted)

        if header_logo is not None:
            try:
                norm, _ = _normalise_image_bytes(header_logo, max_width_px=400)
                para.add_run().add_picture(BytesIO(norm), height=Mm(6.5))
                if text or has_pnum or h_left or h_center or h_right:
                    para.add_run("   ")
            except Exception:
                traceback.print_exc()
        if h_left or h_center or h_right:
            # 3-zone header via tab stops (mirrors the footer layout).
            _dual_tab_stops(para, width_dxa)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if h_left:
                _hrun(h_left)
            para.add_run("\t")
            if h_center:
                _hrun(h_center)
            para.add_run("\t")
            if h_right:
                _hrun(h_right)
            elif has_pnum:
                _add_field(para, "PAGE")
        elif text:
            _hrun(text)
            if has_pnum:
                para.add_run("\t")
                _dual_tab_stops(para, width_dxa)
                _add_field(para, "PAGE")
        elif has_pnum:
            para.add_run("\t")
            _dual_tab_stops(para, width_dxa)
            _add_field(para, "PAGE")
        if (header_logo is not None or text or has_pnum
                or h_left or h_center or h_right):
            _set_paragraph_bottom_border(para, color=theme["rule"], size_pt=4)

    if footer_spec:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0]
        para.text = ""

        left = _smart_quotes((footer_spec.get("left") or "").strip())
        center = _smart_quotes((footer_spec.get("center") or "").strip())
        right = _smart_quotes((footer_spec.get("right") or "").strip())
        text = (footer_spec.get("text") or "").strip()
        has_pnum = bool(footer_spec.get("show_page_numbers"))
        fmt = footer_spec.get("page_number_format") or "{current} / {total}"
        # If the author already embedded a page placeholder in any zone, don't
        # also inject the show_page_numbers fallback (avoids "Page 2 of 3"
        # appearing twice when a template default is merged with a custom zone).
        explicit_pnum = any(
            _has_pagenum_placeholder(z) for z in (left, center, right)
        )
        auto_pnum = has_pnum and not explicit_pnum

        if left or center or right:
            # Explicit 3-zone footer via tab stops.
            _dual_tab_stops(para, width_dxa)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if left:
                if _has_pagenum_placeholder(left):
                    _footer_add_pagenum_runs(para, left, muted)
                else:
                    r = para.add_run(left)
                    r.font.size = Pt(9)
                    r.font.color.rgb = _rgb(muted)
            para.add_run("\t")
            if _has_pagenum_placeholder(center):
                _footer_add_pagenum_runs(para, center, muted)
            elif not center and auto_pnum and not right:
                _footer_add_pagenum_runs(para, fmt, muted)
            elif center:
                r = para.add_run(center)
                r.font.size = Pt(9)
                r.font.color.rgb = _rgb(muted)
            para.add_run("\t")
            if _has_pagenum_placeholder(right):
                _footer_add_pagenum_runs(para, right, muted)
            elif not right and auto_pnum and center:
                _footer_add_pagenum_runs(para, fmt, muted)
            elif right:
                r = para.add_run(right)
                r.font.size = Pt(9)
                r.font.color.rgb = _rgb(muted)
        else:
            # Legacy single-line footer: centered text + page numbers.
            align = (footer_spec.get("alignment") or "center").lower()
            para.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
            if text and _has_pagenum_placeholder(text):
                _footer_add_pagenum_runs(para, _smart_quotes(text), muted)
            elif text:
                run = para.add_run(_smart_quotes(text))
                run.font.size = Pt(9)
                run.font.color.rgb = _rgb(muted)
                if has_pnum:
                    para.add_run("    ")
                    _footer_add_pagenum_runs(para, fmt, muted)
            elif has_pnum:
                _footer_add_pagenum_runs(para, fmt, muted)


def _section_content_height_pt(doc: _DocxDocument) -> float:
    """Usable vertical space of section[0] in points."""
    s = doc.sections[0]
    try:
        h = int(s.page_height) - int(s.top_margin) - int(s.bottom_margin)
    except Exception:
        return 700.0
    return h / 914400.0 * 72.0  # EMU -> pt


async def _cover_image(cover: dict, key: str, resolver) -> Optional[bytes]:
    """Resolve a named image field from a cover dict via the shared pipeline."""
    src = (
        cover.get(key)
        or cover.get(f"{key}_b64")
        or cover.get(f"{key}_image_b64")  # legacy alias (header.logo_image_b64)
        or cover.get(f"{key}_hint")
        or cover.get(f"{key}_url")
    )
    if not src:
        return None
    block: dict = {}
    val = str(src)
    if val.startswith("data:image/") or len(val) > 512:
        block["image_b64"] = val
    elif val.startswith(("http://", "https://")):
        block["image_url"] = val
    elif cover.get(f"{key}_hint") or (key == "logo" and not cover.get(key)):
        block["image_hint"] = cover.get(f"{key}_hint") or val
    else:
        block["image_hint"] = val
    return await resolver(block)


def _cover_hero_image_field(cover: dict) -> bool:
    return bool(
        cover.get("image_b64") or cover.get("image_hint")
        or cover.get("image_generate") or cover.get("image_url")
    )


async def _build_cover_page(
    doc: _DocxDocument,
    cover: dict,
    *,
    accent_hex: str,
    image_resolver,
) -> None:
    """Render a designed cover page.

    Styles (``cover.style``):
      - ``rule`` (default): elegant editorial cover — top accent bar, logo,
        kicker, large title, subtitle, and a footed author/date meta band.
      - ``band``: title reversed out of a full-width accent panel.
      - ``banner``: hero image at the top, title block beneath.
      - ``centered``: simple vertically centred title/subtitle.

    Vertical rhythm is driven by real ``space_before`` values (computed from
    the page height), never by stacks of empty paragraphs.
    """
    theme = _theme(doc)
    style = (cover.get("style") or "rule").lower()
    title = _smart_quotes(cover.get("title") or "")
    subtitle = _smart_quotes(cover.get("subtitle") or "")
    kicker = _smart_quotes(
        cover.get("kicker") or cover.get("eyebrow") or cover.get("doc_type") or ""
    )
    author = _smart_quotes(cover.get("author") or "")
    org = _smart_quotes(cover.get("organization") or cover.get("org") or "")
    date = _smart_quotes(str(cover.get("date") or ""))

    logo = await _cover_image(cover, "logo", image_resolver)
    hero = (
        await image_resolver(cover) if _cover_hero_image_field(cover) else None
    )
    if hero is not None and not cover.get("style"):
        # Author provided a hero image without picking a style -> banner.
        style = "banner"

    content_h = _section_content_height_pt(doc)
    width_dxa = _content_width_dxa(doc) or 9360

    # A full-content-height, borderless 2-row table pins the title block to
    # the top and the meta band to the bottom WITHOUT relying on fragile
    # blank-paragraph spacers (the old approach overflowed onto a 2nd page).
    layout = doc.add_table(rows=2, cols=1)
    layout.autofit = False
    _set_table_width(layout, width_dxa)
    top_cell = layout.rows[0].cells[0]
    meta_cell = layout.rows[1].cells[0]
    top_cell.width = Emu(int(width_dxa * 914400 / 1440))
    meta_cell.width = top_cell.width
    for c in (top_cell, meta_cell):
        _clear_cell_borders(c)
        _set_cell_margins(c, top=0, bottom=0, left=0, right=0)
    top_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    meta_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    _set_row_height(layout.rows[0], content_h * 0.60, exact=False)
    _set_row_height(layout.rows[1], content_h * 0.34, exact=False)

    centered = style == "centered"

    def _new_para(cell, first_used: list):
        if not first_used:
            first_used.append(True)
            return cell.paragraphs[0]
        return cell.add_paragraph()

    used: list = []

    # ── Optional logo (top of the top-cell) ──────────────────────────────
    if logo is not None:
        try:
            norm, _ = _normalise_image_bytes(logo, max_width_px=800)
            lp = _new_para(top_cell, used)
            lp.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
            )
            lp.paragraph_format.space_after = Pt(18)
            lp.add_run().add_picture(BytesIO(norm), height=Mm(14))
        except Exception:
            traceback.print_exc()

    if style == "banner" and hero is not None:
        try:
            norm, _ = _normalise_image_bytes(hero, max_width_px=2000)
            hp = _new_para(top_cell, used)
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hp.paragraph_format.space_after = Pt(22)
            hp.add_run().add_picture(
                BytesIO(norm), width=Mm(width_dxa / 1440 * 25.4),
            )
        except Exception:
            traceback.print_exc()

    if style == "band":
        # Title reversed out of a full-width accent panel (nested 1-cell tbl).
        _render_cover_band_panel(top_cell, doc, title, subtitle, kicker, theme)
    else:
        # Editorial: accent bar + kicker + big title + subtitle.
        if style != "banner":
            bar = _new_para(top_cell, used)
            bar.paragraph_format.space_after = Pt(12)
            if centered:
                bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_bottom_border(bar, color=theme["accent"], size_pt=24)
        if kicker:
            kp = _new_para(top_cell, used)
            if centered:
                kp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            kp.paragraph_format.space_after = Pt(4)
            kr = kp.add_run(kicker.upper())
            kr.bold = True
            kr.font.size = Pt(10.5)
            kr.font.color.rgb = _rgb(theme["accent"])
            _set_run_tracking(kr, 60)
        if title:
            tp = _new_para(top_cell, used)
            if centered:
                tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tp.paragraph_format.space_after = Pt(8)
            tr = tp.add_run(title)
            tr.bold = True
            tr.font.size = Pt(38 if len(title) < 48 else 30)
            tr.font.color.rgb = _rgb(theme["heading"])
        if subtitle:
            sp = _new_para(top_cell, used)
            if centered:
                sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sp.paragraph_format.space_after = Pt(6)
            sr = sp.add_run(subtitle)
            sr.font.size = Pt(15)
            sr.font.color.rgb = _rgb(theme["ink_soft"])

    # ── Footed meta band: author / org / date ────────────────────────────
    meta_parts = [x for x in (author, org) if x]
    if meta_parts or date:
        band = meta_cell.paragraphs[0]
        band.paragraph_format.space_after = Pt(0)
        if centered:
            band.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_top_border(band, color=theme["rule_strong"], size_pt=6)
        band.paragraph_format.space_before = Pt(6)
        # top padding above the rule via an empty run line is unnecessary;
        # the rule sits on top of the band paragraph.
        if meta_parts:
            r = band.add_run("  \u2022  ".join(meta_parts))
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = _rgb(theme["ink"])
        if date:
            if meta_parts and not centered:
                band.add_run("\t")
                band.paragraph_format.tab_stops.add_tab_stop(
                    Emu(int(width_dxa * 914400 / 1440)),
                    WD_TAB_ALIGNMENT.RIGHT,
                )
            elif meta_parts and centered:
                band.add_run("  \u2022  ")
            dr = band.add_run(date)
            dr.font.size = Pt(11)
            dr.font.color.rgb = _rgb(theme["ink_soft"])
    else:
        # Keep the empty meta cell from collapsing weirdly.
        meta_cell.paragraphs[0].add_run("")

    # End of cover: page break onto next page.
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _set_table_width(table, width_dxa: int) -> None:
    """Set an absolute (DXA) table width on the tblPr."""
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_dxa)))
    tbl_w.set(qn("w:type"), "dxa")


def _set_row_height(row, height_pt: float, *, exact: bool = False) -> None:
    """Set a table row height in points (exact or at-least)."""
    try:
        row.height = Pt(max(1.0, height_pt))
        row.height_rule = (
            WD_ROW_HEIGHT_RULE.EXACTLY if exact else WD_ROW_HEIGHT_RULE.AT_LEAST
        )
    except Exception:
        pass


def _render_cover_band_panel(
    cell, doc: _DocxDocument, title: str, subtitle: str, kicker: str, theme: dict,
) -> None:
    """Full-width accent panel (title reversed out) rendered inside a cell."""
    width_dxa = int((_content_width_dxa(doc) or 9360) * 0.98)
    table = cell.add_table(rows=1, cols=1) if hasattr(cell, "add_table") else \
        doc.add_table(rows=1, cols=1)
    table.autofit = False
    _set_table_width(table, width_dxa)
    panel = table.rows[0].cells[0]
    panel.width = Emu(int(width_dxa * 914400 / 1440))
    _set_cell_shading(panel, theme["accent"])
    _clear_cell_borders(panel)
    _set_cell_margins(panel, top=360, bottom=360, left=320, right=320)
    panel.text = ""
    on = theme["on_accent"]
    if kicker:
        kp = panel.paragraphs[0]
        kr = kp.add_run(kicker.upper())
        kr.bold = True
        kr.font.size = Pt(10.5)
        kr.font.color.rgb = _rgb(on)
        _set_run_tracking(kr, 60)
        tp = panel.add_paragraph()
    else:
        tp = panel.paragraphs[0]
    tp.paragraph_format.space_before = Pt(2)
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(34 if len(title) < 48 else 28)
    tr.font.color.rgb = _rgb(on)
    if subtitle:
        sp = panel.add_paragraph()
        sp.paragraph_format.space_before = Pt(6)
        sr = sp.add_run(subtitle)
        sr.font.size = Pt(14)
        sr.font.color.rgb = _rgb(_lighten(theme["accent"], 0.75))


def _set_paragraph_top_border(paragraph, *, color: str, size_pt: int) -> None:
    """Add a thin top border to a paragraph (meta-band separator)."""
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        anchor = None
        for tag in ("w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr"):
            node = p_pr.find(qn(tag))
            if node is not None:
                anchor = node
                break
        if anchor is not None:
            anchor.addprevious(pbdr)
        else:
            p_pr.append(pbdr)
    top = pbdr.find(qn("w:top"))
    if top is None:
        top = OxmlElement("w:top")
        pbdr.insert(0, top)
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size_pt))
    top.set(qn("w:space"), "6")
    top.set(qn("w:color"), _hex_clean(color))


def _clear_cell_borders(cell) -> None:
    """Remove all borders from a table cell (for panels/KPI cards)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = tc_borders.find(qn(f"w:{side}"))
        if b is None:
            b = OxmlElement(f"w:{side}")
            tc_borders.append(b)
        b.set(qn("w:val"), "nil")

# endregion


# region ── Block renderers ────────────────────────────────────────────────────

def _add_inline_runs(paragraph, text: str, *, accent_rgb: tuple[int, int, int]) -> None:
    """Render inline markdown runs into a paragraph.

    Translates the run-spec list from ``_iter_inline_runs`` into Word runs:
    bold/italic/code via run formatting, accent via theme colour, link via
    a hyperlink relationship.
    """
    runs = _iter_inline_runs(text)
    for spec in runs:
        link = spec.get("link")
        if link:
            _add_hyperlink(paragraph, spec.get("text", ""), link)
            continue
        run = paragraph.add_run(spec.get("text", ""))
        if spec.get("bold"):
            run.bold = True
        if spec.get("italic"):
            run.italic = True
        if spec.get("accent"):
            run.font.color.rgb = RGBColor(*accent_rgb)
            run.bold = True
        if spec.get("code"):
            run.font.name = "Consolas"
            try:
                rpr = run._r.get_or_add_rPr()
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = OxmlElement("w:rFonts")
                    rpr.insert(0, rfonts)
                rfonts.set(qn("w:ascii"), "Consolas")
                rfonts.set(qn("w:hAnsi"), "Consolas")
                # Light grey background for inline code
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F2F2F2")
                rpr.append(shd)
            except Exception:
                pass


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Add an external hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    rpr.append(style)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = _smart_quotes(text or url)
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _heading_number(doc: _DocxDocument, level: int) -> str:
    """Return a legal-style number (e.g. ``2.3.1``) for a heading level.

    Counter-based (computed in Python, baked into the run text) so it renders
    identically in Word, LibreOffice and Google Docs without relying on
    field/list-numbering support. Deeper counters reset when a shallower
    heading advances.
    """
    counters = getattr(doc, "_neura_head_counters", None)
    if counters is None:
        counters = {1: 0, 2: 0, 3: 0, 4: 0}
        doc._neura_head_counters = counters  # type: ignore[attr-defined]
    counters[level] += 1
    for deeper in range(level + 1, 5):
        counters[deeper] = 0
    return ".".join(str(counters[lvl]) for lvl in range(1, level + 1))


def _render_heading(doc: _DocxDocument, block: dict, *, accent_rgb) -> None:
    theme = _theme(doc)
    level = max(1, min(4, int(block.get("level") or 1)))
    text = block.get("text") or block.get("title") or ""

    # Optional "eyebrow"/kicker line above the heading (small, tracked, accent).
    eyebrow = block.get("eyebrow") or block.get("kicker")
    if eyebrow:
        ep = doc.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)
        ep.paragraph_format.space_before = Pt(HEADING_BEFORE_PT.get(level, 10))
        ep.paragraph_format.keep_with_next = True
        er = ep.add_run(_smart_quotes(str(eyebrow)).upper())
        er.bold = True
        er.font.size = Pt(8.5)
        er.font.color.rgb = _rgb(theme["accent"])
        _set_run_tracking(er, 40)

    p = doc.add_paragraph(style=f"Heading {level}")
    if eyebrow:
        p.paragraph_format.space_before = Pt(1)
    align = (block.get("align") or "").lower()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Optional automatic section numbering (1, 1.1, 1.1.1). Enabled per-doc
    # via ``styles.numbered_headings`` or per-block via ``numbered``.
    doc_numbering = bool(getattr(doc, "_neura_number_headings", False))
    want_number = block.get("numbered", doc_numbering) and not block.get(
        "no_number", False
    )
    max_level = int(getattr(doc, "_neura_number_max_level", 3) or 3)
    if want_number and level <= max_level:
        # Strip any manual number the author already typed (e.g. "3." or
        # "4.1") so auto-numbering doesn't produce "3  3. Title".
        text = re.sub(r"^\s*\d+(?:\.\d+)*[.)]?\s+", "", text)
        number = _heading_number(doc, level)
        nrun = p.add_run(number + "\u2002\u2002")  # two en-spaces after number
        nrun.bold = True
        # Match the number to the heading tier colour (navy H1 / accent H2+).
        nrun.font.color.rgb = _rgb(theme["heading"] if level == 1 else theme["accent"])
    _add_inline_runs(p, text, accent_rgb=accent_rgb)


def _render_paragraph(doc: _DocxDocument, block: dict, *, accent_rgb) -> None:
    theme = _theme(doc)
    text = block.get("text") or block.get("content") or ""
    if not text.strip():
        return
    p = doc.add_paragraph()
    align = (block.get("align") or "").lower()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    variant = (block.get("variant") or block.get("style") or "").lower()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = DEFAULT_LINE_SPACING
    _add_inline_runs(p, text, accent_rgb=accent_rgb)
    if variant in ("lead", "intro", "standfirst"):
        # A larger, softer intro paragraph that sets up a section.
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.28
        for run in p.runs:
            run.font.size = Pt(theme["size_pt"] + 2)
            if run.font.color.rgb is None:
                run.font.color.rgb = _rgb(theme["ink_soft"])


def _apply_list_numbering(p, num_id: int, level: int) -> None:
    """Attach ``<w:numPr>`` (numId + ilvl) to a paragraph."""
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_pr.append(ilvl)
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    num_pr.append(nid)
    # numPr must precede spacing/ind/jc/rPr in pPr.
    anchor = None
    for tag in ("w:spacing", "w:ind", "w:jc", "w:rPr"):
        node = p_pr.find(qn(tag))
        if node is not None:
            anchor = node
            break
    if anchor is not None:
        anchor.addprevious(num_pr)
    else:
        p_pr.append(num_pr)


def _list_style(block: dict) -> str:
    """Resolve the list style: bullet | decimal | legal | checklist."""
    raw = str(
        block.get("list_style") or block.get("style") or ""
    ).lower().strip()
    if raw in ("legal", "outline", "multilevel"):
        return "legal"
    if raw in ("checklist", "checkbox", "todo", "tasks"):
        return "checklist"
    if raw in ("decimal", "ordered", "numbered", "number"):
        return "decimal"
    if raw in ("bullet", "unordered", "disc"):
        return "bullet"
    # ``ordered`` boolean is the legacy switch.
    ordered = block.get("ordered")
    if isinstance(ordered, str) and ordered.lower() == "legal":
        return "legal"
    return "decimal" if ordered else "bullet"


# Inter-item spacing presets (space AFTER each list item, in points). Compact
# is the enforced default — it reads clean and is what the model should use
# almost always. "relaxed" is an explicit opt-in for short lists on sparse
# pages; it is NEVER auto-triggered (e.g. blank lines between markdown bullets
# do NOT widen the list) so the model can't accidentally produce loose output.
_LIST_SPACING_PT = {
    "tight": 1.0, "compact": 2.0, "snug": 2.0, "normal": 2.0,
    "relaxed": 6.0, "comfortable": 6.0, "spacious": 6.0, "loose": 6.0,
    "airy": 8.0,
}


def _list_spacing_pt(block: dict) -> float:
    raw = str(block.get("spacing") or block.get("density") or "").lower().strip()
    return _LIST_SPACING_PT.get(raw, 2.0)


def _render_list(
    doc: _DocxDocument,
    block: dict,
    *,
    accent_rgb,
    num_ids: dict[str, int],
) -> None:
    theme = _theme(doc)
    items = block.get("items") or block.get("bullets") or []
    if not items:
        return
    style = _list_style(block)

    if style == "checklist":
        _render_checklist(
            doc, items, theme=theme, accent_rgb=accent_rgb,
            after_pt=max(3.0, _list_spacing_pt(block)),
        )
        return

    # Bullet lists share one numId (order-independent). Ordered/legal lists
    # each get a fresh numId so their counter restarts at 1.
    if style == "bullet":
        num_id = num_ids["bullet"]
    else:
        num_id = _mint_list_num(doc, style)

    after_pt = _list_spacing_pt(block)

    def _emit(text: str, level: int) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after_pt)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = DEFAULT_LINE_SPACING
        _apply_list_numbering(p, num_id, level)
        _add_inline_runs(p, text, accent_rgb=accent_rgb)

    def _walk(raw, level: int) -> None:
        level = max(0, min(3, level))
        if isinstance(raw, dict):
            text = raw.get("text") or raw.get("label") or ""
            _emit(str(text), level)
            for sub in (raw.get("items") or raw.get("children") or []):
                _walk(sub, level + 1)
        elif isinstance(raw, (list, tuple)):
            for sub in raw:
                _walk(sub, level + 1)
        else:
            _emit(str(raw), level)

    for raw in items:
        _walk(raw, int(block.get("level") or 0))


def _render_checklist(doc, items, *, theme, accent_rgb, after_pt: float = 3.0) -> None:
    """Render a checklist: leading ☐/☑ glyph + hanging indent, no numbering."""
    for raw in items:
        checked = False
        if isinstance(raw, dict):
            text = raw.get("text") or raw.get("label") or ""
            checked = bool(raw.get("checked") or raw.get("done"))
        else:
            text = str(raw)
            # Recognise a leading marker the model may have inlined.
            stripped = text.lstrip()
            if stripped[:1] in ("\u2610", "\u2611", "\u2612"):
                checked = stripped[0] in ("\u2611", "\u2612")
                text = stripped[1:].lstrip()
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Mm(7)
        pf.first_line_indent = Mm(-7)
        pf.space_after = Pt(after_pt)
        pf.line_spacing = DEFAULT_LINE_SPACING
        gr = p.add_run(("\u2611 " if checked else "\u2610 "))
        gr.font.color.rgb = _rgb(theme["accent"] if checked else theme["ink_soft"])
        _add_inline_runs(p, str(text), accent_rgb=accent_rgb)


def _content_width_dxa(doc: _DocxDocument) -> int:
    """Return the usable content width of section[0] in DXA."""
    s = doc.sections[0]
    width = (s.page_width or 0) - (s.left_margin or 0) - (s.right_margin or 0)
    # python-docx returns Emu objects; .emu attribute or int(width)
    try:
        emu = int(width)
    except Exception:
        emu = 0
    # 1 DXA = 1/1440 inch; 1 EMU = 1/914400 inch → DXA = EMU * 1440 / 914400
    return int(emu * 1440 / 914400)


_NUMERIC_CELL_RE = re.compile(
    r"^[\s]*[+\-\u2212]?[\u20ac$\u00a3]?\s*\d[\d.,\s]*\s*(?:%|pp|k|m|mln|mld|"
    r"€|\$|£|bn|tb|gb|mb)?\s*[\u20ac$\u00a3]?[\s]*$",
    re.IGNORECASE,
)


def _is_numeric_cell(text: str) -> bool:
    t = (text or "").strip()
    if not t or not any(ch.isdigit() for ch in t):
        return False
    return bool(_NUMERIC_CELL_RE.match(t))


def _strip_md_emphasis(text: str) -> str:
    """Strip inline markdown emphasis markers for numeric column sniffing.

    Cells like ``**€ 374.000**`` must still be recognised as numeric so the
    whole column right-aligns; the markers only affect rendering, not type.
    """
    t = (text or "").strip()
    for mark in ("***", "**", "__", "*", "_", "`"):
        t = t.replace(mark, "")
    return t.strip()


def _set_cell_valign(cell, where: str = "center") -> None:
    try:
        cell.vertical_alignment = {
            "top": WD_ALIGN_VERTICAL.TOP,
            "center": WD_ALIGN_VERTICAL.CENTER,
            "bottom": WD_ALIGN_VERTICAL.BOTTOM,
        }.get(where, WD_ALIGN_VERTICAL.CENTER)
    except Exception:
        pass


def _cell_borders_sides(cell, sides: dict) -> None:
    """Set specific borders on a cell. ``sides`` maps side -> (color, size)|None.

    Sides not present are set to ``nil`` (removed) so table-style presets that
    want horizontal-only lines don't inherit stray verticals.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for side in ("top", "left", "bottom", "right"):
        b = tc_borders.find(qn(f"w:{side}"))
        if b is None:
            b = OxmlElement(f"w:{side}")
            tc_borders.append(b)
        spec = sides.get(side)
        if spec is None:
            b.set(qn("w:val"), "nil")
        else:
            color, size = spec
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(size))
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), _hex_clean(color))


def _table_style_name(block: dict) -> str:
    raw = str(block.get("table_style") or block.get("style") or "").lower().strip()
    if raw in ("lines", "line", "horizontal", "rows"):
        return "lines"
    if raw in ("minimal", "clean", "plain"):
        return "minimal"
    if raw in ("grid", "bordered", "box"):
        return "grid"
    return "lines"  # modern default (was implicit "grid" before)


def _next_caption_number(doc: _DocxDocument, kind: str) -> int:
    counters = getattr(doc, "_neura_caption_counters", None)
    if counters is None:
        counters = {}
        doc._neura_caption_counters = counters  # type: ignore[attr-defined]
    counters[kind] = counters.get(kind, 0) + 1
    return counters[kind]


def _render_caption(doc: _DocxDocument, text: str, *, kind: str) -> None:
    """Render a numbered caption (``Table N — ...`` / ``Figure N — ...``)."""
    theme = _theme(doc)
    n = _next_caption_number(doc, kind)
    try:
        p = doc.add_paragraph(style="Caption")
    except Exception:
        p = doc.add_paragraph()
    lbl = p.add_run(f"{kind} {n}")
    lbl.bold = True
    lbl.italic = True
    lbl.font.size = Pt(9)
    lbl.font.color.rgb = _rgb(theme["accent"])
    rest = p.add_run("\u2002\u2014\u2002" + _smart_quotes(str(text)))
    rest.italic = True
    rest.font.size = Pt(9)
    rest.font.color.rgb = _rgb(theme["ink_soft"])


def _render_table(
    doc: _DocxDocument,
    block: dict,
    *,
    accent_rgb,
    accent_hex: str,
) -> None:
    theme = _theme(doc)
    headers = block.get("headers") or block.get("columns") or []
    rows = block.get("rows") or []
    if not headers and not rows:
        return
    num_cols = max(len(headers), max((len(r) for r in rows), default=0))
    if num_cols == 0:
        return
    total_dxa = _content_width_dxa(doc) or 9360

    style = _table_style_name(block)
    zebra = block.get("zebra", style != "minimal")
    total_row = bool(block.get("total_row") or block.get("totals"))
    border_hex = theme["cell_border"]
    header_fill = theme["accent"]
    header_ink = theme["on_accent"]

    # Column widths
    pcts = block.get("column_widths_pct") or [round(100 / num_cols, 4)] * num_cols
    if len(pcts) != num_cols:
        pcts = [round(100 / num_cols, 4)] * num_cols
    col_widths_dxa = [int(total_dxa * p / 100) for p in pcts]
    col_widths_dxa[-1] += total_dxa - sum(col_widths_dxa)

    # Column alignment: explicit ``align`` list (from markdown) wins; else
    # auto-detect numeric columns and right-align them.
    explicit_align = block.get("align") or []
    col_align: list[str] = []
    for ci in range(num_cols):
        a = explicit_align[ci] if ci < len(explicit_align) else ""
        if a in ("left", "center", "right"):
            col_align.append(a)
            continue
        sample = [
            _strip_md_emphasis(str(r[ci]))
            for r in rows if ci < len(r) and str(r[ci]).strip()
        ]
        numeric = sample and all(_is_numeric_cell(s) for s in sample)
        col_align.append("right" if numeric else "left")

    aln_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    table = doc.add_table(rows=0, cols=num_cols)
    table.autofit = False
    _set_table_width(table, total_dxa)

    def _apply_borders(cell, *, is_header: bool, is_last: bool) -> None:
        if style == "grid":
            _set_cell_borders(cell, color=border_hex, size_pt=4)
        elif style == "lines":
            # Horizontal rules only; header carries a stronger accent underline.
            if is_header:
                _cell_borders_sides(cell, {
                    "bottom": (theme["accent"], 12),
                })
            else:
                _cell_borders_sides(cell, {
                    "bottom": (border_hex, 4),
                })
        else:  # minimal
            if is_header:
                _cell_borders_sides(cell, {"bottom": (theme["accent"], 10)})
            else:
                _cell_borders_sides(cell, {})

    # Header row
    if headers:
        row = table.add_row()
        for i in range(num_cols):
            cell = row.cells[i]
            cell.width = Emu(int(col_widths_dxa[i] * 914400 / 1440))
            _set_cell_margins(cell, top=70, bottom=70, left=130, right=130)
            _set_cell_valign(cell, "center")
            # Filled accent header is the default (denser, enterprise look);
            # only ``minimal`` keeps the light text-on-white treatment.
            shade_header = style != "minimal"
            if shade_header:
                _set_cell_shading(cell, header_fill)
            _apply_borders(cell, is_header=True, is_last=False)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = aln_map.get(col_align[i], WD_ALIGN_PARAGRAPH.LEFT)
            h = headers[i] if i < len(headers) else ""
            run = p.add_run(_smart_quotes(str(h)))
            run.bold = True
            run.font.size = Pt(max(9, theme["size_pt"] - 1))
            run.font.color.rgb = _rgb(header_ink if shade_header else theme["accent"])
        # Repeat header across pages
        tr = row._tr
        tr_pr = tr.find(qn("w:trPr"))
        if tr_pr is None:
            tr_pr = OxmlElement("w:trPr")
            tr.insert(0, tr_pr)
        tr_pr.append(OxmlElement("w:tblHeader"))

    # Body rows
    n_body = len(rows)
    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        is_total = total_row and ridx == n_body - 1
        for ci in range(num_cols):
            cell = row.cells[ci]
            cell.width = Emu(int(col_widths_dxa[ci] * 914400 / 1440))
            _set_cell_margins(cell, top=60, bottom=60, left=130, right=130)
            _set_cell_valign(cell, "center")
            _apply_borders(cell, is_header=False, is_last=ridx == n_body - 1)
            if is_total:
                _set_cell_shading(cell, theme["accent_soft"])
                _cell_borders_sides(cell, {
                    "top": (theme["accent"], 8),
                    "bottom": (theme["accent"], 8),
                })
            elif zebra and ridx % 2 == 1:
                _set_cell_shading(cell, theme["zebra"])
            value = row_data[ci] if ci < len(row_data) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = aln_map.get(col_align[ci], WD_ALIGN_PARAGRAPH.LEFT)
            _add_inline_runs(p, str(value), accent_rgb=accent_rgb)
            if is_total:
                for r in p.runs:
                    r.bold = True

    caption = block.get("caption")
    if caption:
        _render_caption(doc, caption, kind=block.get("caption_kind") or "Table")
    else:
        # Breathing room after the table before the next block (Word collapses
        # a bare table against the following paragraph otherwise).
        spacer = doc.add_paragraph()
        pf = spacer.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        spacer.add_run("").font.size = Pt(8)


def _render_kpi_row(doc: _DocxDocument, block: dict, *, accent_rgb) -> None:
    """A row of metric cards (value + label) in a borderless grid.

    ``items``: list of ``{value, label, caption?}`` (or ``{number, ...}``).
    """
    theme = _theme(doc)
    items = block.get("items") or block.get("kpis") or block.get("stats") or []
    items = [i for i in items if isinstance(i, dict)]
    if not items:
        return
    n = min(len(items), 4)
    items = items[:n]
    total_dxa = _content_width_dxa(doc) or 9360
    col = total_dxa // n
    table = doc.add_table(rows=1, cols=n)
    table.autofit = False
    _set_table_width(table, col * n)
    for i, item in enumerate(items):
        cell = table.rows[0].cells[i]
        cell.width = Emu(int(col * 914400 / 1440))
        _clear_cell_borders(cell)
        _set_cell_shading(cell, theme["accent_soft"])
        _set_cell_margins(cell, top=180, bottom=180, left=170, right=170)
        _set_cell_valign(cell, "top")
        cell.text = ""
        value = str(item.get("value") or item.get("number") or item.get("stat") or "")
        label = str(item.get("label") or item.get("title") or item.get("name") or "")
        caption = str(item.get("caption") or item.get("sub") or item.get("note") or "")
        vp = cell.paragraphs[0]
        vp.paragraph_format.space_after = Pt(2)
        vr = vp.add_run(_smart_quotes(value))
        vr.bold = True
        vr.font.size = Pt(24 if len(value) <= 8 else 18)
        vr.font.color.rgb = _rgb(theme["accent"])
        if label:
            lp = cell.add_paragraph()
            lp.paragraph_format.space_after = Pt(0)
            lr = lp.add_run(_smart_quotes(label).upper())
            lr.bold = True
            lr.font.size = Pt(8.5)
            lr.font.color.rgb = _rgb(theme["ink_soft"])
            _set_run_tracking(lr, 30)
        if caption:
            cp = cell.add_paragraph()
            cp.paragraph_format.space_before = Pt(2)
            cr = cp.add_run(_smart_quotes(caption))
            cr.font.size = Pt(8.5)
            cr.font.color.rgb = _rgb(theme["ink_muted"])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _render_definition_list(doc: _DocxDocument, block: dict, *, accent_rgb) -> None:
    """Term/definition pairs: bold accent term, indented description."""
    theme = _theme(doc)
    items = block.get("items") or block.get("definitions") or []
    for raw in items:
        if isinstance(raw, dict):
            term = raw.get("term") or raw.get("title") or raw.get("label") or ""
            desc = (
                raw.get("description") or raw.get("def") or raw.get("text")
                or raw.get("body") or ""
            )
        elif isinstance(raw, (list, tuple)) and len(raw) >= 2:
            term, desc = raw[0], raw[1]
        else:
            continue
        tp = doc.add_paragraph()
        tp.paragraph_format.space_after = Pt(1)
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.keep_with_next = True
        tr = tp.add_run(_smart_quotes(str(term)))
        tr.bold = True
        tr.font.color.rgb = _rgb(theme["accent"])
        dp = doc.add_paragraph()
        dp.paragraph_format.left_indent = Mm(6)
        dp.paragraph_format.space_after = Pt(4)
        _add_inline_runs(dp, str(desc), accent_rgb=accent_rgb)


def _render_code_block(doc: _DocxDocument, block: dict) -> None:
    """Monospace code panel with preserved line breaks and a subtle fill."""
    theme = _theme(doc)
    code = block.get("code") or block.get("text") or block.get("content") or ""
    lang = str(block.get("language") or block.get("lang") or "").strip()
    if not str(code).strip():
        return
    width_dxa = _content_width_dxa(doc) or 9360
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _set_table_width(table, width_dxa)
    cell = table.rows[0].cells[0]
    cell.width = Emu(int(width_dxa * 914400 / 1440))
    _set_cell_shading(cell, "F4F5F7")
    _cell_borders_sides(cell, {"left": (theme["accent"], 18)})
    _set_cell_margins(cell, top=140, bottom=140, left=200, right=160)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.15
    if lang:
        lr = p.add_run(lang.upper() + "\n")
        lr.bold = True
        lr.font.size = Pt(8)
        lr.font.color.rgb = _rgb(theme["ink_muted"])
        _set_run_tracking(lr, 30)
    lines = str(code).replace("\t", "    ").split("\n")
    for idx, line in enumerate(lines):
        run = p.add_run(line if line else "\u00a0")
        run.font.name = theme["mono_font"]
        _force_run_font(run, theme["mono_font"])
        run.font.size = Pt(9.5)
        run.font.color.rgb = _rgb("2B2F36")
        if idx < len(lines) - 1:
            run.add_break()


def _force_run_font(run, font_name: str) -> None:
    """Force a run's ascii/hAnsi/cs font (python-docx only sets ascii by name)."""
    try:
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), font_name)
    except Exception:
        pass


def _render_divider_label(doc: _DocxDocument, block: dict) -> None:
    """Centered small-caps label flanked by a hairline rule."""
    theme = _theme(doc)
    label = block.get("text") or block.get("label") or ""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    if label:
        r = p.add_run(_smart_quotes(str(label)).upper())
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = _rgb(theme["ink_muted"])
        _set_run_tracking(r, 60)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    _set_paragraph_bottom_border(rule, color=theme["rule"], size_pt=6)


def _render_columns(
    doc: _DocxDocument, block: dict, *, accent_rgb, num_ids: dict[str, int],
) -> None:
    """Side-by-side columns via a borderless table.

    ``columns`` (or ``items``): a list of column specs, each a dict with any of
    ``heading``, ``text``, ``bullets``/``items``. 2 or 3 columns supported.
    """
    theme = _theme(doc)
    cols = block.get("columns") or block.get("items") or []
    cols = [c for c in cols if isinstance(c, dict)]
    if not cols:
        return
    n = min(max(len(cols), 1), 3)
    cols = cols[:n]
    total_dxa = _content_width_dxa(doc) or 9360
    col = total_dxa // n
    table = doc.add_table(rows=1, cols=n)
    table.autofit = False
    _set_table_width(table, col * n)
    for i, spec in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.width = Emu(int(col * 914400 / 1440))
        _clear_cell_borders(cell)
        _set_cell_margins(
            cell, top=40, bottom=40,
            left=0 if i == 0 else 150, right=150 if i < n - 1 else 0,
        )
        _set_cell_valign(cell, "top")
        cell.text = ""
        used_first = [False]

        def _para():
            if not used_first[0]:
                used_first[0] = True
                return cell.paragraphs[0]
            return cell.add_paragraph()

        heading = spec.get("heading") or spec.get("title")
        if heading:
            hp = _para()
            hp.paragraph_format.space_after = Pt(3)
            hr = hp.add_run(_smart_quotes(str(heading)))
            hr.bold = True
            hr.font.size = Pt(theme["size_pt"] + 1)
            hr.font.color.rgb = _rgb(theme["heading"])
        text = spec.get("text") or spec.get("body")
        if text:
            for para_text in str(text).split("\n\n"):
                if not para_text.strip():
                    continue
                tp = _para()
                tp.paragraph_format.space_after = Pt(6)
                tp.paragraph_format.line_spacing = DEFAULT_LINE_SPACING
                _add_inline_runs(tp, para_text.strip(), accent_rgb=accent_rgb)
        bullets = spec.get("bullets") or spec.get("list") or []
        for b in bullets:
            bp = _para()
            bp.paragraph_format.left_indent = Mm(5)
            bp.paragraph_format.first_line_indent = Mm(-5)
            bp.paragraph_format.space_after = Pt(2)
            mk = bp.add_run("\u2022  ")
            mk.font.color.rgb = _rgb(theme["accent"])
            _add_inline_runs(
                bp, b if isinstance(b, str) else str(b.get("text", "")),
                accent_rgb=accent_rgb,
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


async def _render_image_block(
    doc: _DocxDocument,
    block: dict,
    *,
    accent_rgb,
    image_resolver,
) -> None:
    theme = _theme(doc)
    raw = await image_resolver(block)
    if not raw:
        return
    normalised, ext = _normalise_image_bytes(raw, max_width_px=1800)
    width_mm = float(block.get("width_mm") or 140.0)
    align = (block.get("align") or "center").lower()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run()
    try:
        run.add_picture(BytesIO(normalised), width=Mm(width_mm))
    except Exception:
        traceback.print_exc()
        return
    caption = block.get("caption")
    if caption:
        # ``figure`` blocks (or numbered_caption) get an auto-numbered
        # "Figure N — ..." caption; plain images keep a simple italic caption.
        if block.get("type") == "figure" or block.get("numbered_caption"):
            _render_caption(doc, caption, kind=block.get("caption_kind") or "Figure")
        else:
            cp = doc.add_paragraph()
            cp.alignment = p.alignment
            crun = cp.add_run(_smart_quotes(str(caption)))
            crun.italic = True
            crun.font.size = Pt(9)
            crun.font.color.rgb = _rgb(theme["ink_soft"])


_CALLOUT_ICONS = {
    "info": "\u2139", "warning": "\u26A0", "note": "\u270E",
    "success": "\u2714", "danger": "\u2716",
}


def _render_callout(doc: _DocxDocument, block: dict, *, accent_rgb) -> None:
    kind = (block.get("kind") or "info").lower()
    palette = CALLOUT_PALETTES.get(kind, CALLOUT_PALETTES["info"])
    icon = _CALLOUT_ICONS.get(kind, _CALLOUT_ICONS["info"])
    border = palette["border"]
    text = block.get("text") or ""
    title = block.get("title") or ""
    width_dxa = _content_width_dxa(doc) or 9360
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _set_table_width(table, width_dxa)
    cell = table.rows[0].cells[0]
    cell.width = Emu(int(width_dxa * 914400 / 1440))
    _set_cell_shading(cell, palette["fill"])
    # Accent bar on the left, hairline elsewhere — cleaner than a full box.
    _cell_borders_sides(cell, {
        "left": (border, 24),
        "top": (palette["fill"], 4),
        "bottom": (palette["fill"], 4),
        "right": (palette["fill"], 4),
    })
    _set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    cell.text = ""
    if title:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        ir = p.add_run(icon + "\u2002")
        ir.bold = True
        ir.font.color.rgb = _rgb(border)
        run = p.add_run(_smart_quotes(title))
        run.bold = True
        run.font.color.rgb = _rgb(border)
        body_p = cell.add_paragraph()
    else:
        body_p = cell.paragraphs[0]
        ir = body_p.add_run(icon + "\u2002")
        ir.bold = True
        ir.font.color.rgb = _rgb(border)
    _add_inline_runs(body_p, text, accent_rgb=accent_rgb)


def _render_quote(doc: _DocxDocument, block: dict, *, accent_rgb) -> None:
    theme = _theme(doc)
    text = block.get("text") or block.get("quote") or ""
    attribution = block.get("attribution") or block.get("author") or ""
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(12)
    p.paragraph_format.right_indent = Mm(8)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.28
    _set_paragraph_left_border(p, color=theme["accent"], size_pt=28)
    run = p.add_run("\u201C" + _smart_quotes(text) + "\u201D")
    run.italic = True
    run.font.size = Pt(theme["size_pt"] + 2)
    run.font.color.rgb = _rgb(theme["ink"])
    if attribution:
        ap = doc.add_paragraph()
        ap.paragraph_format.left_indent = Mm(12)
        ap.paragraph_format.right_indent = Mm(8)
        ap.paragraph_format.space_after = Pt(10)
        _set_paragraph_left_border(ap, color=theme["accent"], size_pt=28)
        arun = ap.add_run("\u2014 " + _smart_quotes(attribution))
        arun.font.size = Pt(10)
        arun.font.color.rgb = _rgb(theme["ink_soft"])


def _set_paragraph_left_border(paragraph, *, color: str, size_pt: int) -> None:
    """Add a thick left border to a paragraph (used for blockquotes)."""
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        p_pr.append(pbdr)
    left = pbdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        pbdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_pt))
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), _hex_clean(color))


def _render_page_break(doc: _DocxDocument) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _render_toc(doc: _DocxDocument, block: dict, *, accent_rgb) -> None:
    title = block.get("title") or "Table of Contents"
    depth = max(1, min(6, int(block.get("depth") or 3)))
    p = doc.add_paragraph(style="Heading 1")
    run = p.add_run(_smart_quotes(title))
    run.bold = True
    toc_p = doc.add_paragraph()
    _add_toc_field(toc_p, depth=depth)
    # Tell Word/LibreOffice to refresh the TOC content when the file is
    # opened. Without this the TOC stays on its placeholder text because
    # fields are cached results, not live formulas.
    _set_update_fields_on_open(doc)
    # Suggest a page break after the TOC (typical for reports)
    if block.get("page_break_after", True):
        _render_page_break(doc)


def _render_signature(doc: _DocxDocument, block: dict, *, accent_rgb) -> None:
    name = block.get("name") or ""
    role = block.get("role") or ""
    date = block.get("date") or ""
    align = (block.get("align") or "left").lower()
    aln = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    # Spacer + thin signature rule
    doc.add_paragraph()
    rule_p = doc.add_paragraph()
    rule_p.alignment = aln
    rule_p.paragraph_format.space_after = Pt(4)
    rule_run = rule_p.add_run("________________________")
    rule_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    if name:
        p = doc.add_paragraph()
        p.alignment = aln
        run = p.add_run(_smart_quotes(name))
        run.bold = True
    if role:
        p = doc.add_paragraph()
        p.alignment = aln
        run = p.add_run(_smart_quotes(role))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    if date:
        p = doc.add_paragraph()
        p.alignment = aln
        run = p.add_run(_smart_quotes(str(date)))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _render_horizontal_rule(doc: _DocxDocument, *, color: str = "CCCCCC") -> None:
    """Thin horizontal rule via paragraph bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    _set_paragraph_bottom_border(p, color=color, size_pt=6)


def _render_memo_header(doc: _DocxDocument, fields: dict, *, accent_rgb) -> None:
    """Render the standard memo block (To:/From:/Date:/Subject:/CC:)."""
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    width_dxa = _content_width_dxa(doc) or 9360
    label_dxa = int(width_dxa * 0.18)
    val_dxa = width_dxa - label_dxa
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    label_aliases = {
        "to": "To",
        "from": "From",
        "date": "Date",
        "subject": "Subject",
        "cc": "CC",
    }
    order = ["to", "from", "date", "subject", "cc"]
    seen = set()
    for key in order + list(fields.keys()):
        if key in seen:
            continue
        if key not in fields:
            continue
        seen.add(key)
        value = str(fields[key]).strip()
        if not value:
            continue
        row = table.add_row()
        c0, c1 = row.cells
        c0.width = Emu(int(label_dxa * 914400 / 1440))
        c1.width = Emu(int(val_dxa * 914400 / 1440))
        _set_cell_margins(c0, top=20, bottom=20, left=0, right=120)
        _set_cell_margins(c1, top=20, bottom=20, left=0, right=0)
        c0.text = ""
        c1.text = ""
        l_run = c0.paragraphs[0].add_run(label_aliases.get(key, key.title()) + ":")
        l_run.bold = True
        l_run.font.color.rgb = RGBColor(*accent_rgb)
        _add_inline_runs(c1.paragraphs[0], value, accent_rgb=accent_rgb)
    # Spacer + rule under memo header
    _render_horizontal_rule(doc, color="888888")


def _render_letter_addresses(doc: _DocxDocument, fields: dict, *, accent_rgb) -> None:
    """Render the typical letter intro: sender, recipient, date, subject."""
    sender = fields.get("sender") or {}
    recipient = fields.get("recipient") or {}
    date = fields.get("date") or ""
    subject = fields.get("subject") or ""

    # Sender (top-left)
    if sender:
        for line in _addr_lines(sender):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(_smart_quotes(line))
            run.font.size = Pt(10)
        doc.add_paragraph()

    # Recipient (right-aligned, optional indent)
    if recipient:
        for line in _addr_lines(recipient):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Mm(95)
            run = p.add_run(_smart_quotes(line))
            run.font.size = Pt(11)

    if date:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(_smart_quotes(str(date)))
        run.italic = True

    if subject:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Subject: ")
        run.bold = True
        sub_run = p.add_run(_smart_quotes(str(subject)))
        sub_run.bold = True


def _addr_lines(addr: dict) -> list[str]:
    """Flatten an address dict into render-ready lines."""
    if isinstance(addr, str):
        return [ln for ln in addr.splitlines() if ln.strip()]
    out: list[str] = []
    for key in ("name", "company", "role", "address", "city", "country", "email", "phone"):
        v = addr.get(key)
        if v:
            out.append(str(v).strip())
    return out

# endregion


# region ── Document orchestrator ─────────────────────────────────────────────

# Block dispatch: a tiny dict-based router. Async types live in ``ASYNC_BLOCKS``
# because they need to await the image pipeline. Everything else is sync.
SYNC_BLOCK_TYPES = {
    "heading", "paragraph", "para", "text", "list", "bullets", "ordered_list",
    "checklist", "definition", "definition_list", "code", "code_block",
    "kpi", "kpi_row", "stat_row", "stats", "columns", "divider_label",
    "table", "callout", "quote", "blockquote", "page_break", "pagebreak",
    "toc", "signature", "rule", "divider", "horizontal_rule", "hr",
    "memo_header",
}
ASYNC_BLOCK_TYPES = {"image", "figure"}


def _normalise_page_breaks(blocks: list) -> list:
    """Drop redundant page breaks that would emit blank pages.

    LLMs frequently open the body with ``\\newpage`` even though the cover
    already breaks to a fresh page, and sometimes stack several breaks in a
    row. We drop any page break before the first real content block and
    collapse consecutive breaks into one.
    """
    cleaned: list = []
    seen_content = False
    for blk in blocks:
        btype = (blk.get("type") or "").lower().strip() if isinstance(blk, dict) else ""
        is_break = btype in ("page_break", "pagebreak")
        if is_break:
            if not seen_content:
                continue  # leading break -> would create a blank page
            if cleaned and (cleaned[-1].get("type") or "").lower().strip() in (
                "page_break", "pagebreak"
            ):
                continue  # collapse consecutive breaks
        else:
            seen_content = True
        cleaned.append(blk)
    return cleaned


async def _render_blocks(
    doc: _DocxDocument,
    blocks: list,
    *,
    spec: dict,
    image_resolver,
) -> None:
    """Iterate blocks and dispatch each to its renderer.

    Unknown block types are silently skipped (defensive — the LLM may invent
    aliases). The renderer mutates ``doc`` in place.
    """
    accent_hex = spec.get("styles", {}).get("accent") or "#1E2761"
    accent_rgb = _hex_to_rgb(accent_hex)
    num_ids = _ensure_numbering_definitions(doc, accent_hex)

    blocks = _normalise_page_breaks(blocks or [])

    for raw_block in blocks:
        if not isinstance(raw_block, dict):
            continue
        btype = (raw_block.get("type") or "paragraph").lower().strip()
        try:
            if btype == "heading":
                _render_heading(doc, raw_block, accent_rgb=accent_rgb)
            elif btype in ("paragraph", "para", "text"):
                _render_paragraph(doc, raw_block, accent_rgb=accent_rgb)
            elif btype in ("list", "bullets"):
                _render_list(
                    doc, raw_block, accent_rgb=accent_rgb, num_ids=num_ids,
                )
            elif btype == "ordered_list":
                _render_list(
                    doc, {**raw_block, "ordered": True},
                    accent_rgb=accent_rgb, num_ids=num_ids,
                )
            elif btype == "checklist":
                _render_list(
                    doc, {**raw_block, "list_style": "checklist"},
                    accent_rgb=accent_rgb, num_ids=num_ids,
                )
            elif btype in ("definition", "definition_list"):
                _render_definition_list(doc, raw_block, accent_rgb=accent_rgb)
            elif btype in ("code", "code_block"):
                _render_code_block(doc, raw_block)
            elif btype in ("kpi", "kpi_row", "stat_row", "stats"):
                _render_kpi_row(doc, raw_block, accent_rgb=accent_rgb)
            elif btype == "columns":
                _render_columns(doc, raw_block, accent_rgb=accent_rgb, num_ids=num_ids)
            elif btype == "divider_label":
                _render_divider_label(doc, raw_block)
            elif btype == "table":
                _render_table(
                    doc, raw_block, accent_rgb=accent_rgb, accent_hex=accent_hex,
                )
            elif btype == "callout":
                _render_callout(doc, raw_block, accent_rgb=accent_rgb)
            elif btype in ("quote", "blockquote"):
                _render_quote(doc, raw_block, accent_rgb=accent_rgb)
            elif btype in ("page_break", "pagebreak"):
                _render_page_break(doc)
            elif btype == "toc":
                _render_toc(doc, raw_block, accent_rgb=accent_rgb)
            elif btype == "signature":
                _render_signature(doc, raw_block, accent_rgb=accent_rgb)
            elif btype in ("rule", "divider", "horizontal_rule", "hr"):
                _render_horizontal_rule(
                    doc, color=raw_block.get("color") or _theme(doc)["rule"],
                )
            elif btype == "memo_header":
                _render_memo_header(
                    doc, raw_block.get("fields") or {}, accent_rgb=accent_rgb,
                )
            elif btype in ("image", "figure"):
                await _render_image_block(
                    doc, {**raw_block, "type": btype}, accent_rgb=accent_rgb,
                    image_resolver=image_resolver,
                )
            else:
                # Unknown block type: try to recover by treating it as a
                # paragraph if it carries a ``text``/``content`` field.
                fallback_text = raw_block.get("text") or raw_block.get("content")
                if fallback_text:
                    _render_paragraph(
                        doc, {"text": fallback_text}, accent_rgb=accent_rgb,
                    )
        except Exception:
            traceback.print_exc()
            # Don't let one bad block kill the whole document
            continue


def _document_title(spec: dict) -> str:
    cover = spec.get("cover")
    cover_title = cover.get("title") if isinstance(cover, dict) else None
    return (
        spec.get("title")
        or cover_title
        or spec.get("subject")
        or "document"
    )


def _tool_success_reply(fname: str, download_url: str) -> str:
    """Build the tool return string after a successful save.

    Verbatim-copy and anti-summary rules live **here** (in the tool
    output), not in the host LLM system prompt, so open-source installs
    stay self-contained and the model sees the link next to the rules.
    """
    return (
        "[TOOL_RESULT]\n\n"
        "OUTPUT_FOR_USER — Your **next assistant message** must be **only** "
        "the text between the dashed lines (`---`) below. Copy it exactly "
        "(including the blank line and the `[filename](url)` markdown link). "
        "Do not add summaries, bullet lists, section outlines, or extra "
        "sentences. Do not wrap in code fences. Do not use HTML `<a>`; keep "
        "the markdown link.\n\n"
        "---\n"
        "Here is the Word document:\n\n"
        f"[{fname}]({download_url})\n"
        "---\n\n"
        "If you output anything outside the dashed block, the user may lose "
        "the clickable download link."
    )


# endregion


# region ── Tools class (OpenWebUI entry point) ────────────────────────────────

class Tools:
    """OpenWebUI tool entry point. Methods marked ``async def`` and decorated
    with a function description become callable tools that the LLM can invoke
    via native function calling."""

    class Valves(BaseModel):
        unsplash_access_key: str = Field(
            default="",
            description=(
                "Unsplash API access key. When set, blocks of type 'image' "
                "with an `image_hint` field fetch a stock photo. Leave empty "
                "to disable Unsplash lookups."
            ),
        )
        image_generation_url: str = Field(
            default="",
            description=(
                "OPTIONAL OpenAI-compatible /v1/images/generations endpoint "
                "used as fallback when the in-process OpenWebUI image router "
                "is not reachable. Leave empty in normal OpenWebUI "
                "deployments — `image_generate` blocks will use the same "
                "image backend as the in-chat 'Generate Image' button."
            ),
        )
        image_generation_api_key: str = Field(
            default="",
            description="Bearer token for image_generation_url fallback (if any).",
        )
        docx_export_dir: str = Field(
            default="/app/backend/data/cache/files",
            description=(
                "Fallback path for .docx export when the OpenWebUI Files API "
                "is unavailable. Same convention as the slides/dashboard tool."
            ),
        )
        emit_status: bool = Field(
            default=True,
            description=(
                "Emit progress status events during generation (visible as "
                "the spinner caption under the message). Disable for tests."
            ),
        )
        default_template: str = Field(
            default="blank",
            description=(
                "Template applied when the model omits the `template` field "
                "in its document spec. One of: blank, letter, report, memo, "
                "proposal, minutes."
            ),
        )

    def __init__(self):
        self.valves = self.Valves()
        self.citation = False
        # Tool descriptors for native function calling (OpenWebUI reads these
        # via getattr(tool_module, '__tools__', None) or by introspection
        # depending on version; the legacy "self.tools = [...]" pattern from
        # openwebui_artifacts.py still works in 0.4+).
        self.tools = [self._tool_descriptor()]

    # ── Public tool entry point ──────────────────────────────────────────────
    async def generate_document(
        self,
        content: str,
        __event_emitter__=None,
        __request__=None,
        __user__=None,
    ) -> str:
        """Generate a Word (.docx) document from a structured spec.

        The ``content`` parameter accepts TWO formats (auto-detected):

        1. **Markdown-with-frontmatter** (preferred, more compact, more
           robust on long documents). YAML frontmatter for
           template/cover/header/footer/styles, body in standard Markdown
           plus Pandoc fenced divs for callout/signature/page-break/toc.
           Inline ``==text==`` maps to accent-colored runs.

           Example:

               ---
               template: report
               title: FY 2024 Report
               cover: auto
               ---

               # Executive summary
               Revenue grew **24%** ...

               ::: callout type="success" title="Note"
               Body...
               :::

        2. **JSON** (legacy, fully supported): an object with the
           ``title/template/page/styles/header/footer/cover/blocks`` shape
           described in the tool descriptor.

        Args:
            content: Either Markdown-with-frontmatter or a JSON string.
                Code fences (```` ```json ````, ```` ```markdown ````,
                ```` ```md ````, ```` ```yaml ````) are stripped.

        Returns:
            A self-contained ``[TOOL_RESULT]`` string: it includes
            ``OUTPUT_FOR_USER`` instructions plus the exact lines to copy
            (markdown link inside ``---`` delimiters). A separate ``message``
            event may also emit the link for live preview.
        """
        try:
            spec_raw = self._parse_content(content)
        except Exception as exc:
            return self._error_reply(
                f"Invalid spec (expected JSON or Markdown with frontmatter): {exc}"
            )

        if not isinstance(spec_raw, dict):
            return self._error_reply(
                "The `content` parameter must be a JSON object or Markdown "
                "with YAML frontmatter."
            )

        # Apply default template when the spec doesn't declare one
        spec_raw.setdefault("template", self.valves.default_template)
        spec = _resolve_template(spec_raw)

        await self._emit_status(__event_emitter__, "Building Word document...", done=False)

        try:
            doc = await self._build_document(
                spec, request=__request__, user_dict=__user__
            )
        except Exception as exc:
            traceback.print_exc()
            return self._error_reply(f"Rendering error: {exc}")

        # Materialise to bytes
        buf = BytesIO()
        try:
            doc.save(buf)
        except Exception as exc:
            traceback.print_exc()
            return self._error_reply(f"Error saving DOCX: {exc}")
        docx_bytes = buf.getvalue()

        await self._emit_status(__event_emitter__, "Saving file...", done=False)
        fname, download_url, save_err = self._save_docx(
            docx_bytes,
            title=_document_title(spec),
            request=__request__,
            user_dict=__user__,
        )

        if not download_url:
            await self._emit_status(
                __event_emitter__,
                "Save failed.",
                done=True,
            )
            return self._error_reply(
                f"Document was built but saving failed ({save_err})."
            )

        # Emit the link ALSO as a live message event so the user sees a
        # clickable link the moment the file is saved (before the model
        # finishes streaming its final reply). This is a "preview" channel.
        await self._emit_link(__event_emitter__, fname, download_url)
        await self._emit_status(__event_emitter__, "Document ready.", done=True)

        # Authoritative copy of the link is ONLY in this return string (the
        # model does not see side-channel emitter events). Instructions for
        # what to say next MUST live here — not in the global LLM system
        # prompt — so open-source deployments stay self-contained.
        return _tool_success_reply(fname, download_url)

    # ── Build pipeline ───────────────────────────────────────────────────────
    async def _build_document(
        self,
        spec: dict,
        *,
        request: Any = None,
        user_dict: Optional[dict] = None,
    ) -> _DocxDocument:
        """Assemble the python-docx Document from the resolved spec."""
        doc = Document()

        # Resolve the design-token palette ONCE and attach it to the doc so
        # every renderer reads colours/fonts from a single coherent source.
        theme = _resolve_theme(spec)
        doc._neura_theme = theme  # type: ignore[attr-defined]

        # Section-numbering preferences (read by _render_heading).
        _styles = spec.get("styles") or {}
        doc._neura_number_headings = bool(  # type: ignore[attr-defined]
            _styles.get("numbered_headings", False)
        )
        doc._neura_number_max_level = int(  # type: ignore[attr-defined]
            _styles.get("numbered_headings_depth", 3) or 3
        )

        _apply_page_setup(doc, spec.get("page") or {})
        _apply_styles(doc, spec.get("styles") or {})

        accent_hex = theme["accent"]
        accent_rgb = _hex_to_rgb(accent_hex)

        # Image resolver bound to current valves
        unsplash_key = (self.valves.unsplash_access_key or "").strip()
        ai_url = (self.valves.image_generation_url or "").strip()
        ai_key = (self.valves.image_generation_api_key or "").strip()

        async def _resolver(block: dict) -> Optional[bytes]:
            return await _resolve_image_bytes(
                block,
                unsplash_key=unsplash_key,
                ai_url=ai_url,
                ai_key=ai_key,
                request=request,
                user_dict=user_dict,
            )

        # Normalize header/footer: authors may pass a plain string (running
        # text) or a dict ({text, logo, alignment, show_page_numbers, ...}).
        # Downstream helpers expect a dict, so coerce strings here.
        header_spec = spec.get("header") or None
        if isinstance(header_spec, str):
            header_spec = {"text": header_spec}
        footer_spec = spec.get("footer") or None
        if isinstance(footer_spec, str):
            footer_spec = {"text": footer_spec}

        # Resolve an optional header logo (b64/url/hint) up front.
        header_logo = None
        if isinstance(header_spec, dict):
            header_logo = await _cover_image(header_spec, "logo", _resolver)

        # Suppress the running header/footer on page 1 when a cover exists.
        cover_present = bool(spec.get("cover"))

        # Header / footer
        _build_header_footer(
            doc,
            header_spec=header_spec,
            footer_spec=footer_spec,
            accent_hex=accent_hex,
            header_logo=header_logo,
            first_page_different=cover_present,
        )

        # Cover page (template "auto" cover OR explicit cover dict). In both
        # cases we backfill missing fields from the top-level spec so authors
        # can put title/subtitle/author/date in the frontmatter root and only
        # use ``cover:`` for style/kicker overrides.
        cover_spec = spec.get("cover")
        if cover_spec == "auto":
            cover_spec = {}
        if isinstance(cover_spec, dict):
            merged_cover = dict(cover_spec)
            for key in ("title", "subtitle", "author", "date",
                        "kicker", "eyebrow", "organization", "org"):
                if not merged_cover.get(key) and spec.get(key):
                    merged_cover[key] = spec.get(key)
            # Default cover style from the styles block (template-provided).
            if not merged_cover.get("style"):
                cstyle = (spec.get("styles") or {}).get("cover_style")
                if cstyle:
                    merged_cover["style"] = cstyle
            has_content = any(
                merged_cover.get(k)
                for k in ("title", "subtitle", "author", "date",
                          "kicker", "eyebrow", "logo", "logo_b64", "logo_hint")
            )
            if has_content:
                await _build_cover_page(
                    doc, merged_cover, accent_hex=accent_hex,
                    image_resolver=_resolver,
                )

        # Template-specific intro blocks (memo header, letter addresses)
        template_name = spec.get("template", "blank")
        if template_name == "memo" and spec.get("memo_fields"):
            _render_memo_header(
                doc, spec["memo_fields"], accent_rgb=accent_rgb,
            )
        elif template_name == "letter" and spec.get("letter_fields"):
            _render_letter_addresses(
                doc, spec["letter_fields"], accent_rgb=accent_rgb,
            )

        # Main body blocks
        await _render_blocks(
            doc, spec.get("blocks") or [], spec=spec, image_resolver=_resolver,
        )

        # Letter-specific outro: closing line + signature placeholder
        if template_name == "letter" and spec.get("letter_fields", {}).get("closing"):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(_smart_quotes(spec["letter_fields"]["closing"]))
            run.italic = True

        return doc

    # ── Input parsing (dual: JSON or Markdown-with-frontmatter) ──────────────
    @staticmethod
    def _parse_content(content) -> Any:
        """Accept JSON, Markdown-with-frontmatter, or a pre-decoded dict.

        Auto-detect:
            - dict in -> dict out (defensive, e.g. tests calling internals)
            - string starting with ``{`` or ``[`` -> JSON
            - anything else -> Markdown (with optional YAML frontmatter)

        Code fences (``` ```json``` / ``` ```markdown``` / ``` ```md``` /
        ``` ```yaml```) are stripped if the model wrapped the payload.
        Markdown parsing is lenient and never raises (failures degrade to
        a paragraph-only spec).
        """
        if isinstance(content, dict):
            return content
        if not isinstance(content, str):
            raise ValueError(
                f"unsupported content type: {type(content).__name__}"
            )
        text = content.strip()
        if text.startswith("```"):
            text = re.sub(r"^```(?:json|markdown|md|yaml)?\s*", "", text)
            text = re.sub(r"\s*```\s*$", "", text)
        text = text.strip()
        head = text.lstrip()
        if head.startswith(("{", "[")):
            return json.loads(text)
        return _parse_markdown(text)

    # ── Save / files API ─────────────────────────────────────────────────────
    def _save_docx(
        self,
        docx_bytes: bytes,
        *,
        title: str,
        request=None,
        user_dict=None,
    ) -> tuple[str, Optional[str], Optional[str]]:
        """Persist the DOCX. Returns (filename, download_url, error).

        ``filename`` is the human-readable label shown in chat. The Files API
        path uploads under that pretty name (it serves via a file id and sets
        Content-Disposition, so spaces/accents are fine). The ``/cache/files``
        fallback, however, is a plain static route that 500s on spaces or
        non-ASCII characters — so there we store and link an **ASCII-safe**
        slug while still showing the pretty label.
        """
        display_name = f"{_human_filename(title)}.docx"

        # --- Primary: OpenWebUI Files API (visible in UI Files panel) ---
        if _HAS_OWUI_FILES and request and user_dict:
            try:
                user_model = Users.get_user_by_id(user_dict["id"])
                if user_model:
                    upload = UploadFile(
                        file=BytesIO(docx_bytes),
                        filename=display_name,
                        headers=Headers({
                            "content-type":
                                "application/vnd.openxmlformats-officedocument."
                                "wordprocessingml.document"
                        }),
                    )
                    file_item = upload_file_handler(
                        request=request,
                        file=upload,
                        metadata={},
                        process=False,
                        user=user_model,
                    )
                    if file_item:
                        file_id = getattr(file_item, "id", None)
                        if file_id:
                            return display_name, f"/api/v1/files/{file_id}/content", None
            except Exception:
                traceback.print_exc()

        # --- Fallback: cache/files (download only, ASCII-safe path) ---
        export_dir = (self.valves.docx_export_dir or "").strip() or "/app/backend/data/cache/files"
        try:
            os.makedirs(export_dir, mode=0o775, exist_ok=True)
            # ASCII slug + short suffix: safe for the static route and unique
            # enough that same-titled docs don't overwrite each other.
            stored = f"{_slugify(title)}_{uuid.uuid4().hex[:6]}.docx"
            filepath = os.path.join(export_dir, stored)
            with open(filepath, "wb") as fh:
                fh.write(docx_bytes)
            if os.path.isfile(filepath) and os.path.getsize(filepath) > 0:
                return display_name, f"/cache/files/{stored}", None
        except Exception as exc:
            return display_name, None, str(exc)
        return display_name, None, "could not save file"

    # ── Response helpers ─────────────────────────────────────────────────────
    @staticmethod
    def _error_reply(msg: str) -> str:
        return (
            "[TOOL_RESULT — use the text below as your final reply, "
            "verbatim, unchanged. Do NOT include this instruction line.]\n\n"
            f"Could not generate the Word document: {msg}"
        )

    # ── Status emit ──────────────────────────────────────────────────────────
    async def _emit_status(self, emitter, description: str, *, done: bool) -> None:
        if not emitter or not self.valves.emit_status:
            return
        try:
            await emitter({
                "type": "status",
                "data": {"description": description, "done": done},
            })
        except Exception:
            pass

    @staticmethod
    async def _emit_link(emitter, fname: str, url: str) -> None:
        """Emit the download link as its own message event.

        Markdown ``[name](url)`` renders as a clickable anchor in every
        OpenWebUI version we ship to. Embedding the link in the tool's
        return value is fragile because (a) the model may strip/rephrase it
        and (b) HTML <a> tags are rendered as raw text unless preceded by a
        raw-HTML block. Emitting it here guarantees the user always sees a
        clickable link, even if the model goes off-script.
        """
        if not emitter:
            return
        try:
            await emitter({
                "type": "message",
                "data": {"content": f"\n\n[{fname}]({url})\n"},
            })
        except Exception:
            pass

    # ── Tool descriptor (OpenWebUI native function calling) ──────────────────
    @staticmethod
    def _tool_descriptor() -> dict:
        return {
            "type": "function",
            "function": {
                "name": "generate_document",
                "description": (
                    "Generate a professional Word (.docx) document and save "
                    "it via the OpenWebUI Files API. On success, the return "
                    "value is a self-contained [TOOL_RESULT] blob that "
                    "includes OUTPUT_FOR_USER instructions and the exact "
                    "markdown download line — follow that blob, not any "
                    "external system prompt.\n\n"
                    "Use when the user asks for: a Word document, .docx, "
                    "business letter, report, memo, internal note, proposal, "
                    "meeting minutes, contract, policy, manual. NEVER "
                    "fabricate a .docx yourself — always use this tool.\n\n"
                    "INPUT (auto-detected — prefer Markdown):\n"
                    "1) Markdown with YAML frontmatter (preferred — fewer "
                    "tokens than JSON, robust on long docs):\n"
                    "   ---\n"
                    "   template: report\n"
                    "   title: FY 2024 Report\n"
                    "   cover: auto\n"
                    "   ---\n"
                    "   # Executive summary\n"
                    "   Revenue grew **24%** with margin up ==6 pp==.\n"
                    "   ::: callout type=\"success\" title=\"Note\"\n"
                    "   Highlight body.\n"
                    "   :::\n"
                    "2) JSON object (legacy, still supported) — see schema "
                    "below.\n\n"
                    "TEMPLATES (set in frontmatter `template:` or JSON "
                    "`template`). Each ships a coherent design system "
                    "(palette derived from `accent`, typographic + spacing "
                    "scale, section rules under H1/H2):\n"
                    "- report — numbered headings (1, 1.1, 1.1.1), editorial "
                    "'rule' cover, footer 'N / M'. Add `[[toc]]` right after "
                    "the cover.\n"
                    "- whitepaper — long-form technical doc; numbered "
                    "headings (depth 2), teal accent, 'rule' cover.\n"
                    "- proposal — commercial proposal; bold accent 'band' "
                    "cover, page numbers, NO heading numbers.\n"
                    "- letter — business letter (30 mm margins, no page "
                    "numbers, no section rules). Use `letter_fields:` or the "
                    "`::: letter` fenced div.\n"
                    "- memo — internal memo. Use `memo_fields:` (to, from, "
                    "date, subject) or the `::: memo` fenced div.\n"
                    "- minutes — meeting minutes (page numbers, green "
                    "accent, section rules).\n"
                    "- blank — base styling only, no default decoration.\n\n"
                    "COVER (frontmatter `cover:` map or JSON `cover`): fields "
                    "`title/subtitle/author/date` (backfilled from the root), "
                    "plus `kicker` (eyebrow label), `organization`, `logo` "
                    "(b64/url/unsplash), and `style`: `rule` (editorial, "
                    "default), `band` (title reversed on an accent panel), "
                    "`banner` (hero image on top), `centered`.\n\n"
                    "HEADER/FOOTER: header supports `text`, `logo` and "
                    "`show_page_numbers`; footer supports either `text` + "
                    "`show_page_numbers`, or a 3-zone layout via `left`, "
                    "`center`, `right` (use the literal `{page}` in a zone "
                    "for the page number). When a cover exists the running "
                    "header/footer is auto-suppressed on page 1.\n\n"
                    "MARKDOWN SYNTAX (body):\n"
                    "- Headings: `# H1`..`#### H4` (auto-numbered on report/"
                    "whitepaper). Add an eyebrow with JSON `eyebrow`.\n"
                    "- Paragraphs: inline `**bold**`, `*italic*`, "
                    "`==accent==`, `` `code` ``, `[link](url)`\n"
                    "- Lists: `-`/`*` (unordered), `1.` (ordered, restarts "
                    "per list), nested via 2-space indent. Rendered COMPACT "
                    "by default (blank lines between items do NOT add "
                    "spacing).\n"
                    "- Checklist: `- [ ] todo` / `- [x] done` (rendered with "
                    "real ☐/☑ glyphs)\n"
                    "- Tables: GFM `| h |\\n|--:|\\n| 12 |` — numeric columns "
                    "auto right-align; header gets an accent underline; JSON "
                    "adds `table_style` (lines|grid|minimal), `total_row`, "
                    "`caption`\n"
                    "- Code: fenced ```lang … ``` renders a monospace panel\n"
                    "- KPI cards: `::: kpi\\nitems:\\n  - value: \\\"30%\\\"\\n"
                    "    label: Saved\\n:::`\n"
                    "- Columns: `::: columns\\ncolumns:\\n  - heading: A\\n"
                    "    text: ...\\n  - heading: B\\n    text: ...\\n:::`\n"
                    "- Definition list: `::: definition\\nitems:\\n  - term: X"
                    "\\n    description: ...\\n:::`\n"
                    "- Figure (numbered caption): `::: figure\\nsrc: "
                    "unsplash:...\\ncaption: ...\\n:::`\n"
                    "- Callouts: `::: callout type=\"info|success|warning"
                    "|danger\" title=\"...\"\\nbody\\n:::` OR GitHub "
                    "admonition `> [!note] Title\\n> body`\n"
                    "- Quote: `> text` — Signature: `::: signature\\nName\\n"
                    "Role\\n2026-04-18\\n:::`\n"
                    "- Page break: `\\\\newpage` OR `::: page-break\\n:::` — "
                    "TOC: `[[toc]]` OR `::: toc depth=\"3\"\\n:::` — "
                    "Horizontal rule: `---`\n\n"
                    "QUALITY RULES (both formats):\n"
                    "- Set the template explicitly — A4 portrait is default.\n"
                    "- Use real content, never placeholders like "
                    "'Lorem ipsum'.\n"
                    "- Reports/whitepapers: a cover + `[[toc]]`, short "
                    "paragraphs, and a mix of headings, lists, tables, KPI "
                    "cards and callouts. Lead a section with a `lead` "
                    "paragraph (JSON `variant:\"lead\"`).\n"
                    "- Tables: <= 6 columns; widths via `column_widths_pct` "
                    "(JSON) or pipe alignment; set `total_row:true` for a "
                    "highlighted totals row.\n"
                    "- List spacing: lists are COMPACT by default — always "
                    "prefer this, it looks clean. Only for a SHORT list that "
                    "must fill a sparse page set JSON `spacing:\"relaxed\"` "
                    "on that one list block. Never make long lists relaxed.\n"
                    "- Images: `unsplash:` for stock (English, scene-based), "
                    "`gen:` for AI (English, 10-30 words), "
                    "`data:image/...;base64,...` for logos.\n"
                    "- `==text==` maps to accent-colored runs."
                ),
                "parameters": {
                    "type": "object",
                    "properties": {
                        "content": {
                            "type": "string",
                            "description": (
                                "EITHER Markdown-with-frontmatter (preferred, "
                                "see description) OR a JSON object with the "
                                "shape:\n"
                                "{\n"
                                "  \"title\": \"...\",\n"
                                "  \"subtitle\": \"...\",\n"
                                "  \"author\": \"...\",\n"
                                "  \"date\": \"YYYY-MM-DD\",\n"
                                "  \"template\": \"blank|letter|report|memo"
                                "|proposal|minutes|whitepaper\",\n"
                                "  \"page\": {\"size\": \"a4\", "
                                "\"orientation\": \"portrait\", "
                                "\"margin_mm\": 25},\n"
                                "  \"styles\": {\"font\": \"Calibri\", "
                                "\"size_pt\": 11, \"accent\": \"#1E2761\", "
                                "\"numbered_headings\": true, "
                                "\"cover_style\": \"rule|band|banner\"},\n"
                                "  \"header\": {\"text\": \"...\", "
                                "\"logo\": \"data:image/png;base64,...\", "
                                "\"show_page_numbers\": false},\n"
                                "  \"footer\": {\"left\": \"Confidential\", "
                                "\"center\": \"{page}\", \"right\": \"Org\"},\n"
                                "  \"cover\": {\"style\": \"rule|band|banner\","
                                " \"kicker\": \"...\", \"title\": \"...\", "
                                "\"subtitle\": \"...\", \"author\": \"...\", "
                                "\"organization\": \"...\", \"date\": \"...\", "
                                "\"logo\": \"data:...\"},\n"
                                "  \"letter_fields\": {\"sender\": {...}, "
                                "\"recipient\": {...}, \"date\": \"...\", "
                                "\"subject\": \"...\", \"closing\": \"...\"},\n"
                                "  \"memo_fields\": {\"to\": \"...\", "
                                "\"from\": \"...\", \"date\": \"...\", "
                                "\"subject\": \"...\"},\n"
                                "  \"blocks\": [\n"
                                "    {\"type\": \"heading\", \"level\": 1, "
                                "\"text\": \"...\", \"eyebrow\": \"...\"},\n"
                                "    {\"type\": \"paragraph\", "
                                "\"text\": \"...\", \"variant\": \"lead\"},\n"
                                "    {\"type\": \"list\", \"list_style\": "
                                "\"bullet|decimal|legal|checklist\", "
                                "\"spacing\": \"compact|relaxed\", "
                                "\"items\": [...]},\n"
                                "    {\"type\": \"table\", \"table_style\": "
                                "\"lines|grid|minimal\", \"total_row\": true, "
                                "\"caption\": \"...\", \"headers\": [...], "
                                "\"rows\": [[...]]},\n"
                                "    {\"type\": \"kpi\", \"items\": "
                                "[{\"value\": \"30%\", \"label\": \"...\"}]},\n"
                                "    {\"type\": \"columns\", \"columns\": "
                                "[{\"heading\": \"...\", \"text\": \"...\"}]},\n"
                                "    {\"type\": \"definition\", \"items\": "
                                "[{\"term\": \"...\", \"description\": "
                                "\"...\"}]},\n"
                                "    {\"type\": \"code\", \"language\": "
                                "\"python\", \"code\": \"...\"},\n"
                                "    {\"type\": \"figure\", \"image_hint\": "
                                "\"...\", \"caption\": \"...\"},\n"
                                "    {\"type\": \"callout\", \"kind\": "
                                "\"info|success|warning|danger\", "
                                "\"title\": \"...\", \"text\": \"...\"}\n"
                                "  ]\n"
                                "}\n\n"
                                "Markdown is preferred for documents > 1 page; "
                                "JSON is preferred only when the spec is "
                                "highly structured (e.g. many tables with "
                                "explicit `column_widths_pct`)."
                            ),
                        },
                    },
                    "required": ["content"],
                },
            },
        }

# endregion


